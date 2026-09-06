# -*- coding: utf-8 -*-
"""
Build script for the MGMT 405 EMBA Hybrid course calendar (.docx).

Source of truth: _calendar_content.py (content + date engine).
Run:  python _build_calendar.py
Output: "Calendar <SECTION> Hybrid -- Fall 2026.docx" in this folder.
Pass --section femba for the FEMBA calendar; the default is EMBA.

Design: navy/gold/cream palette matching the 405 slide decks.
Page 1 = title + semester-at-a-glance agenda with internal links.
Page 2 = before-the-course notes. Then one week per page.
"""

import sys, io, os
import re
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import ImageFont

# --section has to land in the environment BEFORE _calendar_content is
# imported: that module reads MGMT405_SECTION at import time and everything
# below is bound from it (2026-09-05, Nico).
for _i, _a in enumerate(sys.argv):
    if _a == "--section" and _i + 1 < len(sys.argv):
        os.environ["MGMT405_SECTION"] = sys.argv[_i + 1].lower()
    elif _a.startswith("--section="):
        os.environ["MGMT405_SECTION"] = _a.split("=", 1)[1].lower()

from _calendar_content import (ANCHOR_FRIDAY, TERM, LINKS, COURSE_TITLE, SUBTITLE,
                               CALENDAR_NOTE, SYLLABUS_NOTE, TA_NAME, CLASSROOM,
                               CLASS_TIMES, TEXTBOOK_NOTES, MATH_REFRESHER_INTRO,
                               MATH_REFRESHER_ITEMS, SIGNIN_NOTE, WEEKS,
                               WEBSITE_LEAD, WEBSITE_TEXT, inclass_modules,
                               podcast_when, dt, fmt, span,
                               CALENDAR_DOCX, TA_EMAIL, class_when,
                               class_days_line, SITE_BASE, slides_for,
                               slides_pub_name)

# OUT is derived from this script's own folder: the project lives on
# different drive letters on different machines (C:, D:, H:).
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.environ.get("CALENDAR_OUT") or os.path.join(
    HERE, CALENDAR_DOCX + ".docx")

# ---------------- palette ----------------
NAVY  = "0B2B4E"
GOLD  = "E09F3E"
PALEGOLD = "F6E8C9"
CREAM = "FDF6E6"
GRAY  = "555B66"
LIGHT = "C8CDD3"
HDRGRAY = "D9D9D9"     # weekend deadline header
BODYGRAY = "F3F4F6"    # weekend deadline body / agenda deadline rows
AGGRAY = "E9EBEE"
# light-grey body of a podcast card (2026-09-03, Nico): videos are
# yellow and reading is white, so podcasts take a light grey.
PODGRAY = "EFF1F4"
# A problem set is dark red, matching the website (2026-09-03, Nico).
# Word has no alpha, so DUEWASH is C00000 at 7% composited over white.
DARKRED = "C00000"
DUEWASH = "FBEDED"

# Category coding, matching the slide decks (2026-08-28, Nico):
#   in class = dark blue   videos = light yellow   exams = darker yellow
# (was: in class = gold, videos = gray, exams = pale gold)
INCLASS  = NAVY         # in-class rows / card headers
PALEBLUE = "E7EDF4"     # light-blue body of an in-class card
VIDEOYEL = PALEGOLD     # light yellow: video rows / video card bodies
EXAMYEL  = GOLD         # darker yellow: exam rows / exam card header
AGHDRGRAY = GRAY        # dark-gray agenda header row (was navy)
AGHDR_SZ = 12           # agenda header text (was 10.5)
AGHDR_GAP_PT = 8.0      # white spacer row under the agenda header
BLUEGRAY = "9DB0C4"   # thin week separators in the agenda
LINKC = "365F91"
BLACK = "1A1A1A"

CONTENT_W = 6.9   # 8.5 - 2*0.8

# Vertical rhythm of the week pages. EVERY week has to fit on one page
# (2026-08-28, Nico); week 1 is the tightest, so these are snug. Card
# backing shapes are sized from the same numbers via _measure_par, so
# changing one here moves the drawn box with the text.
BULLET_AFTER_PT = 1     # gap under a bullet / note line (was 2)
GLABEL_BEFORE_PT = 4    # gap above an italic group label (was 5)
CARD_GAP_PT = 4         # gap before a weekend / exam / holiday card (was 6)
CONTAINER_AFTER_PT = 6  # clearance under the prep container (was 10)
CARD_PAR_BEFORE_PT = 1  # gap above a rounded card (was 2)
CARD_PAR_AFTER_PT = 2   # gap below a rounded card (was 3)

KIND_META = {
    # agenda fill, band right label builder
    "oncampus":     dict(fill=INCLASS,  legend="On-campus class"),
    "deadline":     dict(fill=VIDEOYEL, legend="Video content"),
    "midterm":      dict(fill=EXAMYEL,  legend="Exam"),
    "final":        dict(fill=EXAMYEL,  legend=None),
    # Weeks 10 and 11 have no in-person component, so the agenda colors them
    # as video content rather than leaving them white (2026-09-03, Nico).
    "thanksgiving": dict(fill=VIDEOYEL, legend=None),
    "examprep":     dict(fill=VIDEOYEL, legend=None),
}

# ---------------- low-level helpers ----------------

def rgb(hexstr):
    return RGBColor.from_string(hexstr)


def style_run(r, text=None, bold=False, italic=False, color=BLACK, size=11,
              underline=False, font="Calibri"):
    if text is not None:
        r.text = text
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = font
    r.font.color.rgb = rgb(color)
    return r


def add_run(p, text, **kw):
    return style_run(p.add_run(), text, **kw)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


# The dark-red problem-set box in the agenda's Due / Exams column
# (2026-09-05, Nico). Narrow enough to sit inside the 1.32" column.
PSET_BOX_W = 1.12
# 8.5 pt + 8 pt of text needs ~0.29"; 0.275" clipped the date's
# descenders (2026-09-05, Nico). 0.32" leaves a hair of padding.
PSET_BOX_H = 0.32


def pset_box(cell, lines):
    """Draw the rounded dark-red box round a problem set AND its due date.

    The first cut used paragraph borders (w:pBdr + w:shd). Those are always
    square and can only wrap one paragraph, and Nico asked for rounded
    corners with the date inside -- so this became a drawn roundRect, the
    same device the week pages use for their problem-set cards.

    Both lines are set in NAVY: the on-campus rows print white on navy, and
    white on the box's pale fill would be unreadable."""
    def pop(cell_, inner_w):
        q = cp(cell_)
        add_run(q, lines[0], bold=True, size=8.5, color=NAVY)
        for extra in lines[1:]:
            q2 = cell_.add_paragraph()
            q2.paragraph_format.space_before = Pt(0)
            q2.paragraph_format.space_after = Pt(0)
            add_run(q2, extra, size=8, color=NAVY)

    # 8.5 pt + 8 pt of text, no padding: page 1 of the calendar is full
    # to the point, and the box has to cost no more than the two plain
    # lines it replaced (2026-09-05).
    return rounded_card(cell, pop, fill=DUEWASH, border=DARKRED,
                        width_in=PSET_BOX_W, compact=True,
                        height_in=PSET_BOX_H)


def cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Each side: (sz_eighths_pt, hexcolor) or 'nil'."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for name, spec in (("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)):
        if spec is None:
            continue
        el = OxmlElement(f'w:{name}')
        if spec == "nil":
            el.set(qn('w:val'), 'nil')
        else:
            sz, col = spec
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), col)
        borders.append(el)
    tcPr.append(borders)


def fixed_table(doc, widths, rows=1):
    t = doc.add_table(rows=rows, cols=len(widths))
    t.autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    # cell margins
    mar = OxmlElement('w:tblCellMar')
    for side, val in (("top", 50), ("bottom", 50), ("left", 110), ("right", 110)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    # keep tblGrid in sync with the cell widths (Word's fixed layout follows
    # the grid; python-docx leaves equal-width gridCols behind)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(w * 1440)))
    return t


def cp(cell):
    """First (empty) paragraph of a cell, cleaned."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_hyperlink(p, url, text, bold=True, italic=False, color=LINKC, size=11,
                  underline=False):
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement('w:hyperlink')
    h.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    h.append(r)
    p._p.append(h)


def add_hyperlink_runs(p, url, parts, bold=True, color=LINKC, size=11):
    """One hyperlink whose text is several runs: [(text, underline), ...].

    The podcast bullets need the timing word underlined INSIDE the link, the
    way the website sets it, and add_hyperlink() writes a single run.
    """
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement('w:hyperlink')
    h.set(qn('r:id'), r_id)
    for text, underline in parts:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(rf)
        if bold:
            rPr.append(OxmlElement('w:b'))
        if underline:
            u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
        col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2)))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        h.append(r)
    p._p.append(h)


def add_runs(p, parts, size=11, **kw):
    """The same run list, unlinked -- an episode that is not uploaded yet."""
    for text, underline in parts:
        add_run(p, text, size=size, underline=underline, **kw)


def podcast_parts(text):
    """"Podcast: Intro to Module 1" -> the runs for "Podcast (before class):
    Intro to Module 1", with the timing word underlined (2026-09-04, Nico --
    the calendar matches the website). The phrase after that word comes from
    podcast_when() in the content module, so a wrap-up can say "after
    watching the Module 3 videos" where the class only did applications
    (2026-09-05). Anything that is not one of the two module episodes is
    left exactly as it is."""
    w = podcast_when(text)
    if w is None:
        return [(text, False)]
    when, tail = w
    rest = text[len("Podcast: "):] if text.startswith("Podcast: ") else text
    return [("Podcast (", False), (when, True), (" %s): " % tail, False),
            (rest, False)]


def add_internal_link(p, anchor, text, bold=True, color=LINKC, size=11):
    h = OxmlElement('w:hyperlink')
    h.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement('w:b'))
    col = OxmlElement('w:color'); col.set(qn('w:val'), color); rPr.append(col)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    h.append(r)
    p._p.append(h)


_bookmark_id = [0]

def add_bookmark(p, name):
    _bookmark_id[0] += 1
    bid = str(_bookmark_id[0])
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bid); start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bid)
    p._p.insert(0, start)
    p._p.append(end)


def gold_rule(doc, sz=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), GOLD)
    bdr.append(bottom)
    pPr.append(bdr)
    return p


def tight_page_break(doc):
    """Page break in a 1pt-tall paragraph so it never spills onto its own page."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)
    r = add_run(p, "", size=1)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    add_run(p, "", size=2)
    return p

# ---------------- rounded category boxes (native shapes) ----------------

_FCACHE = {}

def _mfont(size_pt, bold, italic):
    fname = {(0, 0): "calibri.ttf", (1, 0): "calibrib.ttf",
             (0, 1): "calibrii.ttf", (1, 1): "calibriz.ttf"}[(int(bold), int(italic))]
    key = (fname, int(size_pt * 4))
    if key not in _FCACHE:
        _FCACHE[key] = ImageFont.truetype("C:\\Windows\\Fonts\\" + fname,
                                          int(size_pt * 4))
    return _FCACHE[key]


def _measure_par(p_el, inner_w_pt):
    """Approximate rendered height (pt) of a w:p at the given width."""
    indent_pt = before = after = 0
    pPr = p_el.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None and ind.get(qn('w:left')):
            indent_pt = int(ind.get(qn('w:left'))) / 20
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            before = int(sp.get(qn('w:before')) or 0) / 20
            after = int(sp.get(qn('w:after')) or 0) / 20
    texts, max_sz, bold, italic = [], 9.5, False, False
    for r in p_el.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        sz = 11
        if rPr is not None:
            szel = rPr.find(qn('w:sz'))
            if szel is not None:
                sz = int(szel.get(qn('w:val'))) / 2
            bold = bold or rPr.find(qn('w:b')) is not None
            italic = italic or rPr.find(qn('w:i')) is not None
        max_sz = max(max_sz, sz)
        for tel in r.findall(qn('w:t')):
            texts.append(tel.text or "")
    text = " ".join("".join(texts).replace("\u2022", " ").split())
    if not text:
        return max_sz * 1.26 + before + after
    f = _mfont(max_sz, bold, italic)
    avail = inner_w_pt - indent_pt
    sp_w = f.getlength(" ") / 4
    lines, cur = 1, 0.0
    for wd in text.split(" "):
        w_wd = f.getlength(wd) / 4
        add = w_wd if cur == 0 else sp_w + w_wd
        if cur + add > avail and cur > 0:
            lines += 1
            cur = w_wd
        else:
            cur += add
    return lines * max_sz * 1.26 + before + after


_docpr_id = [1000]

WEEK_CARD_W = CONTENT_W - 2 * (0.5 / 2.54)   # 0.5 cm inset each side vs week band
INNER_CARD_W = WEEK_CARD_W - 0.4             # category cards inside the container


def card_header(cell, text, fill, text_color="FFFFFF", size=11, glyph=None,
                glyph_color=PALEGOLD, inner_w_in=None):
    """Shaded header bar paragraph inside a rounded card."""
    hp = cp(cell)
    hp.paragraph_format.space_after = Pt(4)
    if glyph and inner_w_in:
        hp.paragraph_format.tab_stops.add_tab_stop(Inches(inner_w_in - 0.20),
                                                   WD_TAB_ALIGNMENT.RIGHT)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    hp._p.get_or_add_pPr().append(shd)
    add_run(hp, " " + text, bold=True, color=text_color, size=size)
    if glyph:
        gr = add_run(hp, "\t" + glyph, color=glyph_color, size=13)
        gr.font.name = "Segoe UI Symbol"
    return hp


def rounded_card(doc, populate, fill="FFFFFF", border=NAVY, border_w=9525,
                 width_in=CONTENT_W, compact=False, height_in=None):
    """Rounded-rect text-box shape with drop shadow.

    populate(cell, inner_w_in) renders the card's paragraphs into a temp
    table cell; they are then transplanted into the shape's text box.
    """
    inner_w_in = width_in - 0.24    # minus 0.12" text insets each side
    t = fixed_table(doc, [inner_w_in])
    cell = t.rows[0].cells[0]
    populate(cell, inner_w_in)
    ps = [ch for ch in cell._tc if ch.tag == qn('w:p')]
    total_pt = sum(_measure_par(p, inner_w_in * 72) for p in ps)
    # A compact card has to live inside a table cell (the agenda's Due /
    # Exams column), so it trades the vertical breathing room -- and the drop
    # shadow, whose effectExtent is what reserves layout space below an
    # inline shape -- for fitting the page (2026-09-05).
    ins_v = 9144 if compact else 54864           # EMU: 0.01" vs 0.06"
    eff = ('<wp:effectExtent l="0" t="0" r="0" b="0"/>' if compact else
           '<wp:effectExtent l="0" t="0" r="76200" b="76200"/>')
    shadow = ('' if compact else
              '<a:effectLst><a:outerShdw blurRad="63500" dist="27940" '
              'dir="5400000" rotWithShape="0"><a:srgbClr val="000000">'
              '<a:alpha val="30000"/></a:srgbClr></a:outerShdw></a:effectLst>')
    h_in = height_in if height_in else total_pt / 72 + (0.0 if compact else 0.18)
    # _measure_par floors its font size at 9.5 pt, so a card set in
    # smaller type measures larger than it is -- hence the override.
    cx = int(width_in * 914400)
    cy = int(h_in * 914400)
    adj = int(min(18000, max(3000, 0.12 / h_in * 100000)))
    _docpr_id[0] += 1
    xml = (
        '<w:drawing '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'{eff}'
        f'<wp:docPr id="{_docpr_id[0]}" name="Card {_docpr_id[0]}"/>'
        '<a:graphic><a:graphicData '
        'uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr/>'
        '<wps:spPr bwMode="auto">'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="roundRect"><a:avLst>'
        f'<a:gd name="adj" fmla="val {adj}"/></a:avLst></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>'
        f'<a:ln w="{border_w}"><a:solidFill><a:srgbClr val="{border}"/></a:solidFill></a:ln>'
        f'{shadow}'
        '</wps:spPr>'
        '<wps:txbx><w:txbxContent/></wps:txbx>'
        f'<wps:bodyPr rot="0" vert="horz" wrap="square" lIns="109728" '
        f'tIns="{ins_v}" rIns="109728" bIns="{ins_v}" anchor="t" '
        f'anchorCtr="0"><a:noAutofit/></wps:bodyPr>'
        '</wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing>'
    )
    drawing = parse_xml(xml)
    txbx = drawing.find('.//' + qn('w:txbxContent'))
    for p_el in ps:
        txbx.append(p_el)
    # insert shape paragraph (centered), drop the temp table
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(0 if compact else CARD_PAR_BEFORE_PT)
    par.paragraph_format.space_after = Pt(0 if compact else CARD_PAR_AFTER_PT)
    if compact:
        # An inline drawing sits in a text line, and that line's normal
        # leading is what made the boxed rows ~5 pt taller than the plain
        # text they replaced. Pin the line to the shape's own height.
        par.paragraph_format.line_spacing = Pt(h_in * 72)
    run = par.add_run()
    run._r.append(drawing)
    t._tbl.getparent().remove(t._tbl)
    return h_in


def build_rounded_box(doc, title, glyph, body_fill, groups):
    """Category card: navy header bar + glyph + content groups."""
    def pop(cell, inner_w):
        card_header(cell, title, NAVY, glyph=glyph, inner_w_in=inner_w)
        for grp in groups:
            render_group(cell, grp)
    return rounded_card(doc, pop, fill=body_fill, border=NAVY, width_in=INNER_CARD_W)


def container_box(par, width_in, height_in, ln_w=28575, ln_color=NAVY,
                  fill="alpha0"):
    """Rounded rect floating behind the text, anchored to `par` (extends
    0.06in above it). fill: 'alpha0' (transparent) or a hex color.
    ln_w=0 draws no border. Default = thick navy container."""
    cx = int(width_in * 914400)
    cy = int(height_in * 914400)
    adj = int(min(18000, max(2200, 0.14 / height_in * 100000)))
    if fill == "alpha0":
        # transparent fill via alpha-0 white: a bare <a:noFill/> plus outerShdw
        # hangs Word's PDF export (verified by bisection 2026-08-14)
        fill_xml = ('<a:solidFill><a:srgbClr val="FFFFFF">'
                    '<a:alpha val="0"/></a:srgbClr></a:solidFill>')
    else:
        fill_xml = f'<a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>'
    if ln_w:
        ln_xml = (f'<a:ln w="{ln_w}"><a:solidFill>'
                  f'<a:srgbClr val="{ln_color}"/></a:solidFill></a:ln>')
    else:
        ln_xml = '<a:ln><a:noFill/></a:ln>'
    _docpr_id[0] += 1
    xml = (
        '<w:drawing '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        'relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" '
        'allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:align>center</wp:align></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>-54864</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="76200" b="76200"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{_docpr_id[0]}" name="Container {_docpr_id[0]}"/>'
        '<a:graphic><a:graphicData '
        'uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr/>'
        '<wps:spPr bwMode="auto">'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="roundRect"><a:avLst>'
        f'<a:gd name="adj" fmla="val {adj}"/></a:avLst></a:prstGeom>'
        f'{fill_xml}'
        f'{ln_xml}'
        '<a:effectLst><a:outerShdw blurRad="63500" dist="27940" dir="5400000" '
        'rotWithShape="0"><a:srgbClr val="000000"><a:alpha val="30000"/>'
        '</a:srgbClr></a:outerShdw></a:effectLst>'
        '</wps:spPr>'
        '<wps:bodyPr rot="0" vert="horz" anchor="t"/>'
        '</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>'
    )
    par.add_run()._r.append(parse_xml(xml))


def rounded_picture(doc, path, width_in):
    """Centered inline picture with rounded corners and drop shadow."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    pic = run.add_picture(path, width=Inches(width_in))
    inline = pic._inline
    ee = inline.find(qn('wp:effectExtent'))
    if ee is not None:
        ee.set('r', '76200')
        ee.set('b', '76200')
    spPr = inline.graphic.graphicData.pic.spPr
    prst = spPr.find(qn('a:prstGeom'))
    prst.set('prst', 'roundRect')
    av = OxmlElement('a:avLst')
    gd = OxmlElement('a:gd')
    gd.set('name', 'adj')
    gd.set('fmla', 'val 4000')
    av.append(gd)
    prst.append(av)
    spPr.append(parse_xml(
        '<a:effectLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:outerShdw blurRad="63500" dist="27940" dir="5400000" rotWithShape="0">'
        '<a:srgbClr val="000000"><a:alpha val="30000"/></a:srgbClr>'
        '</a:outerShdw></a:effectLst>'))
    return pic


def render_groups_measured(doc, groups):
    """Render loose groups into the body and return their height (inches)."""
    inner_w_in = CONTENT_W - 0.5
    t = fixed_table(doc, [inner_w_in])
    cell = t.rows[0].cells[0]
    for g in groups:
        render_group(cell, g)
    ps = [ch for ch in cell._tc if ch.tag == qn('w:p')]
    if ps and not ps[0].findall('.//' + qn('w:t')):
        ps = ps[1:]     # drop the cell's initial empty paragraph
    total_pt = sum(_measure_par(p, inner_w_in * 72) for p in ps)
    tbl = t._tbl
    for p_el in ps:
        pPr = p_el.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            cur = int(ind.get(qn('w:left')) or 0)
            ind.set(qn('w:left'), str(cur + 288))   # +0.2in inside container
        tbl.addprevious(p_el)
    tbl.getparent().remove(tbl)
    return total_pt / 72

# ---------------- content renderers ----------------

def render_segments(p, segs, base_size=11):
    for seg in segs:
        if seg[0] == "t":
            add_run(p, seg[1], size=base_size)
        elif seg[0] == "l":
            add_hyperlink(p, LINKS[seg[1]], seg[2], size=base_size)


def bullet_par(container, size=11, indent=0.30):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.first_line_indent = Inches(-0.16)
    pf.tab_stops.add_tab_stop(Inches(indent))
    pf.space_before = Pt(0)
    pf.space_after = Pt(BULLET_AFTER_PT)
    add_run(p, "•\t", color=NAVY, size=size)
    return p


def render_item(container, item, size=11, module=None):
    kind = item[0]
    if kind == "note":
        p = container.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.30)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(BULLET_AFTER_PT)
        add_run(p, item[1], italic=True, color=GRAY, size=size - 0.5)
        return
    p = bullet_par(container, size=size)
    if kind == "t":
        add_run(p, item[1], size=size)
    elif kind == "b":
        add_run(p, item[1], bold=True, color=NAVY, size=size)
    elif kind == "v":
        # ("v", linkkey|None, text, minutes|None). minutes None prints
        # "(++)" -- the video has been re-recorded and not yet measured;
        # linkkey None prints the title as plain text (no Panopto link).
        if item[1]:
            add_hyperlink(p, LINKS[item[1]], item[2], size=size)
        else:
            add_run(p, item[2], size=size)
        # One placeholder, "(link to follow)", and only where the link is
        # genuinely missing. A linked video whose running time is unknown
        # gets no marker -- "(++)" was cryptic (2026-09-06, Nico).
        if not item[1]:
            add_run(p, "  (link to follow)", color=GRAY, size=size - 1,
                    italic=True)
        elif item[3] is not None:
            add_run(p, f"  ({item[3]} min)", color=GRAY, size=size - 1)
        # the slide deck behind this video, published next to the website.
        # Inline, so a seven-video week does not gain seven lines.
        deck = slides_for(module, item[2])
        if deck:
            add_run(p, "  ·  ", color=GRAY, size=size - 1)
            add_hyperlink(p, "%s/slides/%s" % (SITE_BASE, slides_pub_name(deck)),
                          "slides", bold=False, size=size - 1)
    elif kind == "p":
        # module podcast: ("p", url|None, text, minutes|None). The url is
        # a literal Dropbox share link, not a LINKS key; both it and the
        # duration stay absent until the episode has been uploaded.
        parts = podcast_parts(item[2])
        if item[1]:
            add_hyperlink_runs(p, item[1], parts, size=size)
        else:
            add_runs(p, parts, size=size)
        if item[3]:
            add_run(p, f"  ({item[3]} min)", color=GRAY, size=size - 1)
    elif kind == "l":
        add_hyperlink(p, LINKS[item[1]], item[2], size=size)
    elif kind == "mix":
        render_segments(p, item[1], base_size=size)


def render_group(container, group, size=11, budget=True):
    if group.get("label"):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(GLABEL_BEFORE_PT)
        p.paragraph_format.space_after = Pt(BULLET_AFTER_PT)
        add_run(p, group["label"], italic=True, color=NAVY, size=size)
    # A video title does not always name its module ("Video 2: The
    # Production Function"), so fall back to the group's label.
    label_mods = re.findall(r"Module (\d+)", group.get("label") or "")
    for item in group["items"]:
        own = re.findall(r"Module (\d+)", item[2] if len(item) > 2
                         and isinstance(item[2], str) else "")
        mods = own or label_mods
        render_item(container, item, size=size,
                    module=int(mods[0]) if mods else None)
    if budget:
        vids = [i for i in group["items"] if i[0] == "v"]
        known = [i[3] for i in vids if i[3] is not None]
        mins = sum(known)
        n_vid = len(vids)
        # a group with any unmeasured video has no meaningful total
        if n_vid >= 2 and mins >= 20 and len(known) == n_vid:
            p = container.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.30)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(BULLET_AFTER_PT)
            add_run(p, f"≈ {mins} min of video in total",
                    italic=True, color=GRAY, size=9.5)

# ---------------- document setup ----------------

def setup_document():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    pf = normal.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.different_first_page_header_footer = True

    # footer: "MGMT 405 · Fall 2026" left, page number right
    fp = sec.footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_W), WD_TAB_ALIGNMENT.RIGHT)
    add_run(fp, f"MGMT 405 · {TERM}\t", color=GRAY, size=9)
    for fld, txt in (("begin", None), (None, " PAGE "), ("end", None)):
        r = fp.add_run()
        style_run(r, "", color=GRAY, size=9)
        if fld:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), fld)
            r._r.append(el)
        else:
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve')
            el.text = txt
            r._r.append(el)
    return doc

# ---------------- page 1: title + agenda ----------------

def week_span_dates(wk):
    if "span_override" in wk:
        (w1, d1), (w2, d2) = wk["span_override"]
        return dt(w1, d1), dt(w2, d2)
    n = wk["num"]
    return dt(n, "Mon"), dt(n, "Sun")


def exam_window(wk):
    """(start, end) dates of an exam week's exam, read off its own window so
    p.1, the week band and the exam box always agree (2026-08-31)."""
    (wd0, off0), (wd1, off1) = wk["exam"]["window"]
    return dt(wk["num"] + off0, wd0), dt(wk["num"] + off1, wd1)


def agenda_due_text(wk):
    """(text, is_pset) per entry. The name is spelled out in full -- it was
    abbreviated to "PS X" until 2026-09-05, when Nico asked for the full
    "Problem Set X" inside a dark-red box."""
    parts = []
    for label, w, d, note in wk["due"]:
        pset = label.startswith("Problem Set")
        if w:
            parts.append((f"{label}\n{fmt(dt(w, d), wd=True)}", pset))
        else:
            parts.append((label, pset))
    if wk["kind"] == "midterm":
        a, b = exam_window(wk)
        parts.insert(0, (f"Midterm window\n{span(a, b)}", False))
    if wk["kind"] == "final":
        a, b = exam_window(wk)
        parts.insert(0, (f"Final Exam window\n{span(a, b)}", False))
    return parts


def build_page1(doc):
    p = doc.add_paragraph()
    add_run(p, COURSE_TITLE, bold=True, color=NAVY, size=20)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, SUBTITLE, bold=True, color=NAVY, size=13)
    gold_rule(doc, sz=20, space_after=5)

    # info block
    for label, value in (("TA:", None), ("Classroom:", CLASSROOM),
                         ("Class times:", CLASS_TIMES)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        # 12pt hand-tweaked by Nico on 2026-08-15 (was 10.5)
        add_run(p, f"{label}  ", bold=True, color=NAVY, size=12)
        if label == "TA:":
            # FEMBA's mailbox is not set yet; print the name alone rather
            # than route students to the other section's TA inbox
            if TA_EMAIL:
                add_run(p, f"{TA_NAME} (", size=12)
                add_hyperlink(p, LINKS["ta_email"], TA_EMAIL,
                              bold=False, size=12)
                add_run(p, ")", size=12)
            else:
                add_run(p, f"{TA_NAME}  (email address to follow)", size=12)
        else:
            add_run(p, value, size=12)

    spacer(doc, 2)

    # Course-website callout (cream rounded card)
    def pop_bruin(cell, inner_w):
        # The course website leads the callout (2026-09-04, Nico): it is
        # where the current calendar, the syllabus and every video live.
        p = cp(cell)
        add_run(p, WEBSITE_LEAD + "  ", bold=True, color=NAVY, size=12)
        add_hyperlink(p, LINKS["website"], WEBSITE_TEXT, size=12,
                      underline=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        add_run(p, CALENDAR_NOTE, bold=True, color=NAVY, size=11)
        # wording + own line hand-tweaked by Nico on 2026-08-15
        p2 = cell.add_paragraph()
        add_run(p2, "Time stamp for this version: Last updated "
                    f"{datetime.now().strftime('%B %d, %Y')}.",
                italic=True, color=GRAY, size=9.5)
    rounded_card(doc, pop_bruin, fill=CREAM, border=NAVY)

    spacer(doc, 3)

    # anchor paragraph for the rounded backing card (card holds heading,
    # note, table, and legend)
    anchor_par = doc.add_paragraph()
    anchor_par.paragraph_format.space_before = Pt(0)
    anchor_par.paragraph_format.space_after = Pt(0)
    anchor_par.paragraph_format.line_spacing = Pt(1)
    add_run(anchor_par, "", size=1)

    # agenda heading (inside the card)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Quarter at a Glance", bold=True, color=NAVY, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Click a week to jump to its page.", italic=True, color=GRAY, size=9.5)
    head_h = (5 + 2 + 14 * 1.26 + 4 + 9.5 * 1.26) / 72

    widths = [0.82, 1.58, 2.98, 1.32]        # 6.70 total, inset on the backing
    # +1 header row, +1 white spacer row separating it from Week 1
    t = fixed_table(doc, widths, rows=2 + len(WEEKS))
    # The agenda's rows carry a shade less vertical padding than a
    # default table. Page 1 is full to the point -- heading, table and
    # legend all have to sit inside ONE drawn backing card -- and the
    # problem-set boxes in the Due column need the room (2026-09-05).
    _cellmar = t._tbl.tblPr.find(qn('w:tblCellMar'))
    for _side in ('top', 'bottom'):
        _cellmar.find(qn('w:' + _side)).set(qn('w:w'), '30')
    # explicit width + centered, with w:jc in proper schema position
    # (appending w:jc after w:tblLook makes Word's placement unreliable)
    tblW = t._tbl.tblPr.find(qn('w:tblW'))
    tblW.set(qn('w:w'), str(int(sum(widths) * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblW.addnext(jc)
    # tighter vertical cell padding than the default helper (agenda only)
    mar = t._tbl.tblPr.find(qn('w:tblCellMar'))
    for side in ('top', 'bottom'):
        mar.find(qn(f'w:{side}')).set(qn('w:w'), '20')
    hdr = t.rows[0]
    for i, htxt in enumerate(("Week", "Dates", "Topics Covered", "Due / Exams")):
        c = hdr.cells[i]
        shade_cell(c, AGHDRGRAY)
        add_run(cp(c), htxt, bold=True, color="FFFFFF", size=AGHDR_SZ)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # white gap row, so the dark-gray header reads apart from Week 1's
    # dark-blue row instead of merging into one dark block
    gap = t.rows[1]
    for c in gap.cells:
        shade_cell(c, "FFFFFF")
        cell_borders(c, top="nil", bottom="nil", left="nil", right="nil")
        gp = cp(c)
        gp.paragraph_format.space_before = Pt(0)
        gp.paragraph_format.space_after = Pt(0)
        gp.paragraph_format.line_spacing = Pt(AGHDR_GAP_PT)
        add_run(gp, "", size=1)
    gtr = gap._tr.get_or_add_trPr()
    gth = OxmlElement('w:trHeight')
    gth.set(qn('w:val'), str(int(AGHDR_GAP_PT * 20)))
    gth.set(qn('w:hRule'), 'exact')
    gtr.append(gth)

    last_ri = len(WEEKS) + 1
    for ri, wk in enumerate(WEEKS, start=2):
        row = t.rows[ri]
        kind = wk["kind"]
        fill = KIND_META[kind]["fill"]
        # text colors follow the category fill: white on the dark-blue
        # in-class rows, navy on the darker-yellow exam rows, black else
        if kind == "oncampus":
            maintxt, subtxt, linkc = "FFFFFF", PALEGOLD, "FFFFFF"
        elif kind in ("midterm", "final"):
            maintxt, subtxt, linkc = NAVY, NAVY, NAVY
        else:
            maintxt, subtxt, linkc = BLACK, GRAY, LINKC
        d1, d2 = week_span_dates(wk)
        cells = row.cells
        for c in cells:
            if fill != "FFFFFF":
                shade_cell(c, fill)
            # thin blue-gray rule after each week (not after the last row)
            if ri < last_ri:
                cell_borders(c, bottom=(4, BLUEGRAY))
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Week link
        p = cp(cells[0])
        add_internal_link(p, f"Week{wk['num']}", f"Week {wk['num']}", size=10.5,
                          color=linkc)
        # Dates
        p = cp(cells[1])
        add_run(p, span(d1, d2), size=10.5, color=maintxt,
                bold=(kind in ("oncampus", "final")))
        sub = None
        if kind == "oncampus":
            f, s = dt(wk["num"], "Fri"), dt(wk["num"], "Sat")
            sub = class_days_line(wk["num"])
        if sub:
            p2 = cells[1].add_paragraph()
            add_run(p2, sub, italic=True, size=8.5, color=subtxt)
        # Topics Covered (original notation)
        p = cp(cells[2])
        add_run(p, "; ".join(wk["topics"]), size=10.5,
                bold=(kind in ("midterm", "final")), color=maintxt)
        # Due
        parts = agenda_due_text(wk)
        blank = cells[3].paragraphs[0]      # the cell's own empty paragraph
        used_blank = False
        for part, is_pset in parts:
            lines = part.split("\n")
            if is_pset:
                # a drawn card, appended to the cell -- it brings its own
                # paragraph, so it never touches `blank`
                pset_box(cells[3], lines)
                continue
            if not used_blank:
                p = cp(cells[3])
                used_blank = True
            else:
                p = cells[3].add_paragraph()
            add_run(p, lines[0], bold=True, size=9.5,
                    color=("FFFFFF" if kind == "oncampus" else NAVY))
            for extra in lines[1:]:
                p2 = cells[3].add_paragraph()
                add_run(p2, extra, size=9, color=subtxt)
        if parts and not used_blank:
            # every entry was a drawn card; drop the leftover empty paragraph
            # or it prints a blank line above the box
            blank._p.getparent().remove(blank._p)

    # measure rows; force all data rows to the SAME height (the tallest one)
    hdr_h = 0.0
    data_hs = []
    for idx, row in enumerate(t.rows):
        row_h = 0.0
        for ci, c in enumerate(row.cells):
            cell_pt = sum(_measure_par(pe, (widths[ci] - 0.22) * 72)
                          for pe in c._tc if pe.tag == qn('w:p'))
            row_h = max(row_h, cell_pt)
        if idx == 0:
            hdr_h = row_h
        elif idx > 1:            # idx 1 is the fixed-height white gap row
            data_hs.append(row_h)
    # uniform height = typical (median) row; rows with longer topic lists
    # (weeks 5, 9) keep growing via the atLeast rule
    # median + 2 pt of slack. It was + 4 until 2026-09-04, when page 1 grew
    # a line for the course-website link and the legend fell onto page 2 --
    # the backing card is one drawn rectangle, so the whole block has to
    # stay on page 1. 2 pt per row over 12 rows buys the line back and the
    # rows still fit their content (the rule is atLeast, not exact).
    # 2 pt of slack until 2026-09-05; see the padding note above
    uniform_pt = sorted(data_hs)[len(data_hs) // 2]
    for row in t.rows[2:]:
        trPr = row._tr.get_or_add_trPr()
        th = OxmlElement('w:trHeight')
        th.set(qn('w:val'), str(int(uniform_pt * 20)))
        th.set(qn('w:hRule'), 'atLeast')
        trPr.append(th)

    # legend (inside the card, below the table)
    # hairline paragraph so Word does NOT merge the legend table into the
    # agenda table (adjacent tables merge without one)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(1)
    sep.paragraph_format.line_spacing = Pt(1)
    add_run(sep, "", size=1)
    lwidths = [0.28, 1.10, 0.28, 1.70, 0.28, 0.50]
    lt = fixed_table(doc, lwidths)
    ltW = lt._tbl.tblPr.find(qn('w:tblW'))
    ltW.set(qn('w:w'), str(int(sum(lwidths) * 1440)))
    ltW.set(qn('w:type'), 'dxa')
    ljc = OxmlElement('w:jc')
    ljc.set(qn('w:val'), 'center')
    ltW.addnext(ljc)
    lrow = lt.rows[0]
    legend = [(INCLASS, "On-campus class"),
              (VIDEOYEL, "Video content"),
              (EXAMYEL, "Exam")]
    for i, (fill, label) in enumerate(legend):
        sw = lrow.cells[2 * i]
        shade_cell(sw, fill)
        cell_borders(sw, top=(2, LIGHT), bottom=(2, LIGHT),
                     left=(2, LIGHT), right=(2, LIGHT))
        add_run(cp(lrow.cells[2 * i + 1]), " " + label, color=GRAY, size=9)
    legend_h = (9 * 1.26 + 8) / 72

    # white rounded, shadowed backing card behind heading + table + legend
    # (Word renders rows ~7% taller than the twip setting)
    table_h = (hdr_h + 8 + AGHDR_GAP_PT) / 72 + sum(
        max(uniform_pt, h + 4) for h in data_hs) * 1.07 / 72
    container_box(anchor_par, CONTENT_W,
                  head_h + table_h + legend_h + 0.05 + 0.10,
                  ln_w=0, fill="FFFFFF")

# ---------------- page 2: before the course ----------------

def build_page2(doc):
    tight_page_break(doc)
    p = doc.add_paragraph()
    add_run(p, "Before the Course Starts", bold=True, color=NAVY, size=16)
    gold_rule(doc, sz=14, space_after=8)

    # pointer to the full syllabus (cream rounded card)
    def pop_syll(cell, inner_w):
        p = cp(cell)
        add_run(p, SYLLABUS_NOTE, bold=True, color=NAVY, size=11)
    rounded_card(doc, pop_syll, fill=CREAM, border=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "General notes on the textbook:", bold=True, color=NAVY, size=11.5)
    for note in TEXTBOOK_NOTES:
        render_item(doc, ("t", note))

    spacer(doc, 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Optional pre-class training:  ", bold=True, color=NAVY, size=11.5)
    add_run(p, MATH_REFRESHER_INTRO, size=11)
    for segs in MATH_REFRESHER_ITEMS:
        render_item(doc, ("mix", segs))

    spacer(doc, 10)
    # sign-in callout (cream rounded card, centered)
    # 14pt text + 6.835in box width hand-tweaked by Nico on 2026-08-15
    def pop_signin(cell, inner_w):
        p = cp(cell)
        add_run(p, "Watching the course videos:  ", bold=True, color=NAVY, size=14)
        for seg in SIGNIN_NOTE:
            if seg[0] == "t":
                add_run(p, seg[1], size=14)
            else:
                add_hyperlink(p, LINKS[seg[1]], seg[2], size=14, underline=True)
    rounded_card(doc, pop_signin, fill=CREAM, border=NAVY,
                 width_in=6249806 / 914400)

    # Panopto sign-in screenshot, centered, rounded + shadow
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Screenshot:", bold=True, italic=True, color=NAVY, size=11)
    rounded_picture(doc, os.path.join(HERE,
                                      "Images", "Panopto-Login-Picture.png"),
                    width_in=4.4)

# ---------------- week pages ----------------

BAND_W = (4.3, 2.6)        # week band: left cell, right label cell
BAND_LBL_PT = 9.5          # right label size


def _check_band_label(txt):
    """The band's right label has to stay on ONE line -- a wrap makes that
    week's band taller than every other page's. Measure it in the real
    font and fail loudly rather than shipping the wrap."""
    usable = BAND_W[1] - 2 * (110 / 1440.0)   # fixed_table cell margins
    w_in = _mfont(BAND_LBL_PT, True, False).getlength(txt) / 4 / 72.0
    if w_in > usable:
        raise SystemExit("week band label wraps (%.3f in > %.3f in): %r"
                         % (w_in, usable, txt))
    return txt


def band_right_label(wk):
    n = wk["num"]
    k = wk["kind"]
    if k == "oncampus":
        return f"On-campus class: Fri, {fmt(dt(n, 'Fri'))} & Sat, {fmt(dt(n, 'Sat'))}"
    if k == "deadline":
        return (f"Video content  \u00b7  suggested: "
                f"Sun, {fmt(dt(n, 'Sun'))}")
    if k == "midterm":
        return f"Midterm window: {span(*exam_window(wk))}"
    if k == "thanksgiving":
        return "Thanksgiving week"
    if k == "examprep":
        return "Exam preparation"
    if k == "final":
        return f"Final Exam window: {span(*exam_window(wk))}"
    return ""


def build_week(doc, wk):
    n = wk["num"]
    tight_page_break(doc)

    d1, d2 = week_span_dates(wk)

    # navy band with gold underline
    t = fixed_table(doc, list(BAND_W))
    row = t.rows[0]
    for c in row.cells:
        shade_cell(c, NAVY)
        cell_borders(c, bottom=(14, GOLD), top=(4, NAVY),
                     left=(4, NAVY), right=(4, NAVY))
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cp(row.cells[0])
    add_bookmark(p, f"Week{n}")
    add_run(p, f"Week {n}", bold=True, color="FFFFFF", size=15)
    add_run(p, f"   ·   {span(d1, d2)}", bold=True, color=PALEGOLD, size=13)
    p = cp(row.cells[1])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, _check_band_label(band_right_label(wk)),
            bold=True, color="FFFFFF", size=BAND_LBL_PT)

    spacer(doc, 2)

    # due cards at the top, right under the week band
    for label, w, d, note in wk["due"]:
        # Only a problem set takes the dark-red treatment; the practice
        # final is not one, so it keeps the gold rule (2026-09-03, Nico).
        is_pset = label.lower().startswith("problem set")

        def pop_due(cell, inner_w, label=label, w=w, d=d, note=note,
                    is_pset=is_pset):
            p = cp(cell)
            if w:
                add_run(p, f"Due:  {label} – {fmt(dt(w, d), wd=True)}",
                        bold=True, color=NAVY, size=11.5)
            else:
                add_run(p, f"{label} – {note}", bold=True, color=NAVY, size=11.5)
            # Every problem set says where the solution goes, as on the
            # website (2026-09-04, Nico).
            if is_pset:
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_before = Pt(1)
                add_run(p2, "Upload one solution per group on ", size=10)
                add_hyperlink(p2, LINKS["bruinlearn_course"], "BruinLearn",
                              size=10, bold=False, underline=True)
        rounded_card(doc, pop_due,
                     fill=DUEWASH if is_pset else "FFFFFF",
                     border=DARKRED if is_pset else GOLD,
                     border_w=15875, width_in=WEEK_CARD_W)

    # topics card -- on an exam week the exam card takes its place, so the
    # week opens with one dark-yellow-headed box instead of three (2026-08-31)
    if wk.get("exam"):
        ex = wk["exam"]
        (wd0, off0), (wd1, off1) = ex["window"]
        w0 = fmt(dt(n + off0, wd0), wd=True)
        w1 = fmt(dt(n + off1, wd1), wd=True)

        def pop_exam(cell, inner_w):
            card_header(cell, ex["title"], EXAMYEL, text_color=NAVY, size=12)
            for line in ex["lines"]:
                render_item(cell, ("t", line.format(w0=w0, w1=w1)))
        rounded_card(doc, pop_exam, fill=CREAM, border=GOLD,
                     width_in=WEEK_CARD_W)

    def pop_topics(cell, inner_w):
        p = cp(cell)
        add_run(p, "Topics covered", bold=True, color=NAVY, size=11.5)
        for ti, topic in enumerate(wk["topics"]):
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(2 if ti == 0 else 1)
            add_run(p2, topic, size=11)
    if not wk.get("exam"):
        rounded_card(doc, pop_topics, fill="FFFFFF", border=NAVY,
                     width_in=WEEK_CARD_W)

    def render_weekend():
        """The on-campus class card (or, on a video week, the
        "videos to watch" card)."""
        spacer(doc, CARD_GAP_PT)
        we = wk["weekend"]
        da, db = we["days"]
        if wk["kind"] == "oncampus":
            title = f"On-campus class   ·   {class_when(n)}"
            hdr_fill, fill, border = INCLASS, PALEBLUE, NAVY
            # classical building, matching the website (2026-09-03, Nico).
            # The header text measures 5.14" against a 6.07" right tab stop,
            # so the glyph has room and cannot push the line to wrap.
            hdr_glyph = "\U0001F3DB\uFE0F"
        else:
            title = (f"Videos to watch   ·   suggested deadline: "
                     f"{fmt(dt(n, da), wd=True)} / {fmt(dt(n, db), wd=True)}")
            hdr_fill, fill, border = INCLASS, VIDEOYEL, NAVY
            hdr_glyph = None

        def pop_weekend(cell, inner_w):
            card_header(cell, title, hdr_fill, text_color="FFFFFF", size=11.5,
                        glyph=hdr_glyph, inner_w_in=inner_w)
            for g in we["groups"]:
                render_group(cell, g)
            # In-Class Material, as on the website (2026-09-04, Nico): a
            # handout and a slide deck per module the class covers, "(TBD)"
            # until they are uploaded right before class. One line per
            # module -- the two rows the website uses would cost the week
            # pages 3 more lines each, and every week has to stay on one
            # page.
            if wk["kind"] == "oncampus":
                mods = inclass_modules(wk)
                if mods:
                    mlist = ("Module %s" % mods[0] if len(mods) == 1 else
                             "Modules " + ", ".join(str(n) for n in mods))
                    render_group(cell, {
                        "label": f"In-Class Material – {mlist}: "
                                 f"Handout / Slides  (TBD)",
                        "items": []}, budget=False)
        rounded_card(doc, pop_weekend, fill=fill, border=border,
                     width_in=WEEK_CARD_W)

    # An on-campus week LEADS with the class card (2026-09-03, Nico): the
    # class is the anchor of the week, so it sits above the preparation
    # container rather than under it.
    weekend_first = bool(wk.get("weekend")) and wk["kind"] == "oncampus"
    if weekend_first:
        render_weekend()

    # prep section: Videos / Podcasts / Suggested Reading cards inside a
    # transparent thick-navy container box
    if wk["prep_groups"]:
        heading_par = None
        total_h = 0.0
        if wk.get("prep_days"):
            a, b = wk["prep_days"]
            heading_par = doc.add_paragraph()
            heading_par.paragraph_format.space_before = Pt(6)
            heading_par.paragraph_format.space_after = Pt(4)
            heading_par.paragraph_format.left_indent = Inches(0.4)
            label = "Before class" if wk["kind"] == "oncampus" else "During the week"
            add_run(heading_par,
                    f"{label}  ·  {fmt(dt(n, a), wd=True)} – {fmt(dt(n, b), wd=True)}",
                    bold=True, color=NAVY, size=12)
            total_h += (6 + 4 + 12 * 1.26) / 72
        cats, others = {}, []
        for g in wk["prep_groups"]:
            c = g.get("cat", "other")
            (others if c == "other" else cats.setdefault(c, [])).append(g)
        for key, title, glyph, fill in (
                # clapperboard -- the Hollywood "action" symbol -- for the
                # card header (2026-09-03, Nico). It replaced the play
                # triangle and then the film reel. Video BULLETS are
                # unaffected; they carry no glyph of their own.
                ("video", "Videos", "\U0001F3AC", VIDEOYEL),
                ("podcast", "Podcasts", "\U0001F3A7", PODGRAY),
                ("read", "Suggested Reading", "\U0001F4D6", "FFFFFF"),
                ("practice", "Suggested Additional Practice Exercises",
                 "\u270E", "FFFFFF")):
            if key not in cats:
                continue
            h = build_rounded_box(doc, title, glyph, fill, cats[key])
            # card + its paragraph spacing + 3 pt slack
            total_h += h + (CARD_PAR_BEFORE_PT + CARD_PAR_AFTER_PT + 3) / 72
        if others:
            total_h += render_groups_measured(doc, others)
        if heading_par is not None:
            total_h += 0.06 + 0.10          # top offset + bottom padding
            container_box(heading_par, WEEK_CARD_W, total_h)
            spacer(doc, CONTAINER_AFTER_PT)  # clear the container's bottom edge

    # weekend / exam / holiday blocks
    if wk.get("weekend") and not weekend_first:
        render_weekend()

    if wk.get("holiday"):
        spacer(doc, CARD_GAP_PT)
        ho = wk["holiday"]
        (wd0, off0), (wd1, off1) = ho["window"]

        def pop_holiday(cell, inner_w):
            p = cp(cell)
            add_run(p, f"{ho['text']}   ({fmt(dt(n + off0, wd0), wd=True)} – "
                       f"{fmt(dt(n + off1, wd1), wd=True)})",
                    bold=True, color=GRAY, size=11)
        rounded_card(doc, pop_holiday, fill=BODYGRAY, border=LIGHT,
                     width_in=WEEK_CARD_W)


# ---------------- main ----------------

def main():
    doc = setup_document()
    build_page1(doc)
    build_page2(doc)
    for wk in WEEKS:
        build_week(doc, wk)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
