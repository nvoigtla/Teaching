# -*- coding: utf-8 -*-
"""
Build script for the MGMT 405 EMBA Hybrid course SYLLABUS (.docx).

    python _build_syllabus.py            # write the .docx
    python _build_syllabus.py --md       # write the Markdown draft too

Design: the same navy / gold / cream palette and chrome as the course
calendar, so the syllabus, the calendar and the website read as one family.
The layout helpers are IMPORTED from ../Course Calendar/_build_calendar.py
rather than copied -- that module does all its work inside main(), so
importing it is safe -- and every address comes from the calendar's LINKS
registry, so no link is typed twice.

CONTENT is the single source of truth below: SECTIONS drives both the .docx
and the Markdown draft, so the two cannot drift apart.

Two standing content decisions (2026-09-04, Nico):
  * NO e-mail addresses anywhere in this document. The PDF is published on
    the public course website, and the site deliberately obfuscates both
    addresses against harvesters; printing them in a public PDF would undo
    that. The document points at the website's "Class and Contact" box and
    at Bruin Learn instead.
  * Achieve is gone. Practice exercises are the TA's own site, which is
    what the calendar and the website link.
"""

import io
import os
import sys
import argparse

CAL = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                   os.pardir, "Course Calendar"))
if CAL not in sys.path:
    sys.path.insert(0, CAL)

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# --section has to land in the environment BEFORE _calendar_content is
# imported (it reads MGMT405_SECTION at import time), and _build_calendar
# imports it in turn (2026-09-05, Nico).
for _i, _a in enumerate(sys.argv):
    if _a == "--section" and _i + 1 < len(sys.argv):
        os.environ["MGMT405_SECTION"] = sys.argv[_i + 1].lower()
    elif _a.startswith("--section="):
        os.environ["MGMT405_SECTION"] = _a.split("=", 1)[1].lower()

from _build_calendar import (                      # noqa: E402
    NAVY, GOLD, CREAM, GRAY, LIGHT, PALEGOLD, BLACK, CONTENT_W,
    add_run, add_hyperlink, cp, fixed_table, shade_cell, cell_borders,
    gold_rule, spacer, rounded_card, setup_document, render_segments)
import _calendar_content as C                      # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, C.SYLLABUS_DOCX + ".docx")
OUT_MD = os.path.join(HERE, C.SYLLABUS_DOCX + ".md")

TITLE = C.COURSE_TITLE
SUBTITLE = (f"Course Syllabus – {C.TERM} – "
            f"{C.SECTIONS[C.SECTION]['subtitle_tail']}")

# The meeting pattern, spelled out for the prose. EMBA meets twice over a
# weekend, FEMBA once on the Saturday, so the heading changes with it.
_MEETING_DAYS = [C.dt(1, wd).strftime("%A") for wd, _ in C.MEETINGS]
FIRST_MEETING_DAY = _MEETING_DAYS[0]
MEETING_SENTENCE = " and ".join(
    "%ss %s" % (day, time)
    for day, (_, time) in zip(_MEETING_DAYS, C.MEETINGS))
MEETING_HEADING = ("On-campus weekends" if len(C.MEETINGS) > 1
                   else "On-campus sessions")

# ============================== CONTENT ==============================
# Block kinds:
#   ("h",   text)              section heading (navy, gold rule under)
#   ("sub", text)             bold navy sub-heading
#   ("p",   text)             body paragraph
#   ("mix", [segments])       body paragraph with links (calendar segments)
#   ("b",   text)             bullet
#   ("bmix",[segments])       bullet with links
#   ("card",[blocks])         cream rounded callout holding p / mix lines
#   ("grades", rows)          the small grade-weight table
#   ("gap", pts)              vertical air

INSTRUCTOR = "Prof. Nico Voigtländer"
OFFICE = "Anderson C-511, Entrepreneur’s Hall"

# Where the two e-mail addresses now live, said once and reused.
WHERE_MAIL = ("The e-mail addresses are on the course website, in the "
              "“Class and Contact” box, and on Bruin Learn.")

SECTIONS = [
    ("h", "Course Purpose"),
    ("p", "Managerial Economics is concerned with the application of economic "
          "principles to key management decisions. It provides guidance to "
          "increase value creation, and allows a better understanding of the "
          "external business environment in which organizations operate. A "
          "primary purpose of the course is to develop tools useful in other "
          "Anderson courses: economics is a key foundation for much of what is "
          "taught in finance, marketing, business strategy and virtually every "
          "other course in the MBA program. Managerial Economics is a unique "
          "way of thinking about problems and decisions that managers face in "
          "each of the functional areas of their organization. Economics "
          "stresses the importance of incentives as determinants of human "
          "behavior and performance, and emphasizes the consideration of costs "
          "and benefits as an efficient method for reaching economic "
          "decisions."),

    ("h", "How This Hybrid Course Runs"),
    ("p", "The course covers eight modules over twelve weeks. Most of the "
          "material is taught in short videos that you watch before the "
          "weekend; three on-campus weekends then apply it to cases, examples "
          "and exercises. Modules 5 and 8 are taught in class rather than on "
          "video."),
    ("blmix", [("t", "Videos: posted on Panopto and linked from the course "
                    "website. Sign in with the option "),
              ("l", "panopto_site", "“ASM Panopto”"),
              ("t", ". Each module’s videos have a suggested deadline – "
                    "the Sunday of the week, or the %s class on an "
                    "on-campus weekend." % FIRST_MEETING_DAY)]),
    ("bl", "%s: %s, in room %s. They fall in weeks 1, 5 and 9 "
           "– see the course calendar for the dates."
           % (MEETING_HEADING, MEETING_SENTENCE, C.CLASSROOM)),
    ("bl", "Podcasts: each module has two short audio episodes – an intro to "
          "listen to before the module, and a wrap-up to listen to after it."),
    ("blmix", [("t", "Practice: optional online exercises for every module, "
                    "with hints and step-by-step solutions, on the "),
              ("l", "practice_index", "practice-exercise site"),
              ("t", ".")]),
    ("bl", "Assessment: five group problem sets, a midterm covering the "
          "material through Module 3, and a final exam covering all eight "
          "modules."),
    ("p", "The course website carries the whole schedule week by week, and "
          "the course calendar prints the same information for those who "
          "prefer it on paper."),

    ("h", "Textbook"),
    ("p", "The required textbook for the course is Microeconomics, by "
          "Goolsbee, Levitt and Syverson, 4th edition."),
    ("b", "The “Figure it out” boxes in each chapter should be "
          "considered as advanced (voluntary) readings."),
    ("sub", "Optional additional reading"),
    ("p", "For those of you that would like a suggestion of a more basic text "
          "to serve as a foundation, I suggest:"),
    ("b", "The Economic Way of Thinking by Paul Heyne, Peter J. Boettke and "
          "David L. Prychitko, Prentice Hall, 2005."),
    ("b", "Microeconomics: A Very Short Introduction, by Avinash Dixit, "
          "Oxford University Press, 2014."),
    ("sub", "Newspaper articles"),
    ("p", "Articles from the Wall Street Journal or other news sources will "
          "be assigned. These selected articles will be posted on Bruin Learn "
          "a couple of days in advance. These readings will be discussed in "
          "class and everyone will be expected to have read these assigned "
          "articles."),
    ("mix", [("t", "You can subscribe to the WSJ at a reduced educational "
                   "rate (optional – not required for the class): "),
             ("l", "wsj_student", "wsj.com/studentoffer")]),

    ("h", "Online Practice Exercises"),
    ("mix", [("t", "Interactive practice exercises for each module, with "
                   "hints and step-by-step solutions, are available on the "
                   "course’s practice site: "),
             ("l", "practice_index", "rafaelrubiao.github.io/mgmt405-practice"),
             ("t", ". They are also linked from every module page of the "
                   "course website.")]),
    ("p", "These exercises are not mandatory. They are for your own practice "
          "only and do not affect your grade."),

    ("h", "Optional Pre-Class Training: Math Refresher"),
    ("p", "For those of you who feel they would need a math refresher before "
          "the class, I suggest the following. We have designed a math quiz "
          "that helps you to assess whether you are already comfortable with "
          "key math concepts needed in Managerial Economics, or whether you "
          "would benefit from further math preparation. The quiz consists of "
          "10 questions and should not take more than 10 – 15 minutes."),
    ("bmix", [("t", "Take the "), ("l", "math_quiz", "Math Quiz"),
              ("t", " (10 – 15 min)")]),
    ("bmix", [("t", "If you do not score high on the quiz, watch the "),
              ("l", "math_videos", "Math Review Videos"),
              ("t", ", which go over the most important math concepts for "
                    "our class")]),

    ("h", "Contact Information and Office Hours"),
    ("p", f"My office is located in {OFFICE}."),
    ("p", WHERE_MAIL),
    ("p", "Important – how to schedule a meeting: my office hours are by "
          "appointment and the time is flexible. To schedule a (Zoom) "
          "meeting, send me an e-mail with your question and we then "
          "schedule a time."),
    ("b", "Please address all questions about problem sets, practice exams "
          "and practice problems to the TA."),
    ("b", "Address questions about concepts covered in class to me."),
    # "regular", not "weekly" -- the sessions are not necessarily weekly
    # (2026-09-04, Nico).
    ("p", "I will hold regular “Coffee & Econ Sessions” via Zoom. "
          "Times will be announced in class."),

    ("h", "Teaching Assistant"),
    ("p", f"The teaching assistant for this section is {C.TA_NAME}. There is "
          "a section-specific TA e-mail address for Section 2; it is on the "
          "course website and on Bruin Learn."),
    ("p", "The TA will hold a weekly office hour and review session via "
          "Zoom, at a time to be announced. During each session, the TA will "
          "provide supplemental instruction on more technical content, solve "
          "additional exercises, and answer any questions related to course "
          "content or the problem sets. In addition to the weekly sessions, "
          "the TA will conduct review sessions prior to the final exam."),

    ("h", "Grades"),
    ("p", "Final grades are determined on the following basis:"),
    ("grades", [("Midterm Exam", "35%"),
                ("Final Exam", "40%"),
                ("Problem Sets", "25%")]),
    ("p", "It is imperative that you come to class prepared – having read "
          "the assignments beforehand."),
    ("p", "The grade distributions will correspond to the School’s "
          "guidelines for core courses: no more than 20% A+’s and "
          "A’s, and no more than 45% A’s of all types. Under no "
          "circumstances will special assignments substitute for inadequate "
          "performance, nor will extra-credit projects be assigned."),

    ("h", "Midterm and Final Exams"),
    ("p", "Both exams take place online, and you will have a 3.5-hour window "
          "to solve the exam and upload your solutions. The midterm exam "
          "covers the material through Module 3. The final exam covers all "
          "the class content, Modules 1 – 8, and has about 20 multiple "
          "choice questions and 3 – 4 problem-solving questions. To "
          "provide flexibility, you will have the option to take each exam "
          "any time within a given time window, as indicated on the course "
          "calendar; the exact window will be announced in class. No makeup "
          "exams will be scheduled."),
    ("p", "The exams are open-book. You may use the course textbook, class "
          "slides, class notes and problem sets. You may use a calculator or "
          "Excel. You may access the course materials on Bruin Learn and the "
          "online version of the textbook. Usage of the internet for any "
          "other purpose is prohibited. In particular, the use of AI tools "
          "(e.g., ChatGPT) is not allowed during the exams, and both exams "
          "are proctored by an online proctoring company. Communicating with "
          "anyone regarding the exam is prohibited."),
    ("p", "A practice final exam is scheduled in the exam-preparation week, "
          "and a sample final exam with solutions is available on Bruin Learn "
          "as an example of the type of exam given in the past."),

    ("h", "Group Problem Sets"),
    ("p", "There will be 5 problem sets. These problems are designed to check "
          "your progress and to extend and reinforce concepts covered in "
          "class. One or two questions similar to the ones on the problem "
          "sets will appear on the exams. Since the exams test individual "
          "performance, it is critical that you acquire skills at solving the "
          "problems independently. Therefore, while these assignments are to "
          "be turned in by groups, you should independently attempt "
          "beforehand to answer each question. Group assignments will be "
          "graded on a continuous 0 – 100 scale. Solutions will be "
          "distributed for each assignment. These step-by-step solutions are "
          "a good substitute for seeing me solve problems in person in class: "
          "there is usually not enough class time to do this."),
    ("mix", [("t", "Each study group submits one set of answers. The problem "
                   "sets are due according to the schedule in the course "
                   "calendar. Upload one solution per group on "),
             ("l", "bruinlearn_course", "Bruin Learn"),
             ("t", ". No late assignments can be accepted.")]),

    ("h", "Attendance Policy"),
    ("p", "Whether to attend class is your personal decision. There is no "
          "penalty in terms of participation. However, I highly recommend "
          "that you attend each class."),
    ("sub", "If you have to miss a class"),
    ("b", "You do not need to contact the instructor or the TA"),
    ("b", "Do the readings assigned for the week and watch the class "
          "recording"),
    ("b", "Make sure you coordinate work on the problem set with your study "
          "group"),

    ("h", "Use of Unauthorized Materials"),
    ("p", "Do not use any old answer keys to exams and homework sets "
          "distributed in prior courses, unless I have distributed them "
          "myself. Any published work that is used in your written solutions "
          "must contain the appropriate citation. Violations of academic "
          "integrity standards are taken very seriously at UCLA."),
    ("mix", [("t", "UCLA’s Academic Integrity Policy is found here: "),
             ("l", "integrity", "deanofstudents.ucla.edu/Academic-Integrity")]),
    ("sub", "Artificial Intelligence (AI) tools"),
    ("p", "The use of AI tools (e.g. ChatGPT) is prohibited during the "
          "midterm and the final exam. For problem sets, you may use AI tools "
          "to help brainstorm or to revise existing work you have written. "
          "When you submit your problem set, we expect you to clearly "
          "attribute what text was generated by the AI tool (e.g., "
          "AI-generated text appears in a different colored font, quoted "
          "directly in the text, or use an in-text parenthetical citation). "
          "While AI tools can provide helpful insights, we want you to think "
          "critically about the information you receive. We encourage you to "
          "ask specific questions, provide context, and evaluate the quality "
          "of the answers provided by AI tools. Remember, the answers are not "
          "always correct."),

    ("h", "Regrade Policy"),
    ("p", "For regrade requests, submit the original copy of the graded "
          "exam or assignment and a separate page with a written explanation "
          "of the request to the TA. If you cannot come to an agreement with "
          "the TA, I will step in as a mediator. Note that I reserve the "
          "right to regrade the entire exam or assignment. This means that a "
          "regrade request could result in a higher or lower final score."),

    ("h", "For On-Campus Classes: Use of Computers and Tablets"),
    ("p", "No use of laptops or tablets, with the exception of note-taking. "
          "Surfing the internet and catching up on your e-mail is "
          "distracting, both for yourself and for those around you. "
          "Exception: if you prefer to take notes on your tablet – but "
          "for note-taking only."),

    ("h", "Course Materials: Website and Bruin Learn"),
    ("mix", [("t", "The course website, "),
             ("l", "website", C.WEBSITE_TEXT),
             ("t", ", is the fastest way to find what a given week asks of "
                   "you: the videos, the podcasts, the readings, the practice "
                   "exercises and every deadline, week by week and module by "
                   "module. The class syllabus and the course calendar can be "
                   "downloaded there as PDFs.")]),
    ("mix", [("t", "All course materials, including electronic copies of all "
                   "our slides, problem sets and solutions, are on the "),
             ("l", "bruinlearn_course", "Bruin Learn"),
             ("t", " site for the course. This course is almost entirely "
                   "paperless, except for the occasional class handout. "
                   "Please monitor the course pages regularly, as they carry "
                   "the most up-to-date information on the reading "
                   "assignments. I will make extensive use of e-mail to "
                   "communicate with the class, so please check your e-mail "
                   "regularly.")]),
    ("p", "In addition, an Excel workbook used in class is available on the "
          "class site."),

    ("h", "UCLA Center for Accessible Education (CAE)"),
    ("mix", [("t", "The UCLA Center for Accessible Education (CAE) "
                   "facilitates academic accommodations for regularly "
                   "enrolled, matriculating students with documented "
                   "permanent and temporary disabilities. Accommodations are "
                   "designed to promote successful engagement in the UCLA "
                   "academic experience. Students needing academic "
                   "accommodations based on a disability should contact the "
                   "Center for Accessible Education by using the "
                   "“Contact Us Form” at "),
             ("l", "cae", "cae.ucla.edu"),
             ("t", " or at (310) 825-1501. When possible, students should "
                   "contact the CAE within the first two weeks of the term, "
                   "as reasonable notice is needed to coordinate "
                   "accommodations.")]),

    ("h", "Instructor Evaluation"),
    ("p", "At the midpoint of the class, your views about the course will be "
          "formally solicited using an online survey. There will also be a "
          "final course evaluation. You are free to contact me and discuss "
          "the course with me and offer suggestions for improvements at any "
          "time. I am always glad to receive feedback. In addition, an "
          "opportunity for providing me with anonymous feedback exists on "
          "Bruin Learn."),
]

# Addresses the syllabus needs that the calendar's registry did not carry.
EXTRA_LINKS = {
    "practice_index": "https://rafaelrubiao.github.io/mgmt405-practice/index.html",
    "wsj_student": "http://www.wsj.com/studentoffer",
    "integrity": "https://www.deanofstudents.ucla.edu/Academic-Integrity",
    "cae": "https://www.cae.ucla.edu",
}
C.LINKS.update({k: v for k, v in EXTRA_LINKS.items() if k not in C.LINKS})

# ============================== .docx renderer ==============================

BODY_SZ = 11
H_SZ = 13.5


def heading(doc, text):
    spacer(doc, 4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, bold=True, color=NAVY, size=H_SZ)
    gold_rule(doc, sz=8, space_after=4)


def sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, bold=True, color=NAVY, size=11.5)


def body(doc, text=None, segs=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if segs is not None:
        render_segments(p, segs, base_size=BODY_SZ)
    else:
        add_run(p, text, size=BODY_SZ)


def _lead_label(text, lead):
    """"Videos: posted on ..." -> ("Videos:", " posted on ..."), so the
    bullets of "How This Hybrid Course Runs" can be scanned down the left
    edge. Only bullets marked with the "lead" flag are split -- a heuristic
    on the first colon would also bold half of a book title ("Microeconomics:
    A Very Short Introduction")."""
    if not lead:
        return None, text
    i = text.find(":")
    if i <= 0:
        raise ValueError("lead bullet without a label: %r" % (text,))
    return text[:i + 1], text[i + 1:]


def bullet(doc, text=None, segs=None, lead=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.30)
    pf.first_line_indent = Inches(-0.16)
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    add_run(p, "•  ", color=NAVY, size=BODY_SZ)
    if segs is not None:
        segs = list(segs)
        if segs and segs[0][0] == "t":
            lab, rest = _lead_label(segs[0][1], lead)
            if lab:
                add_run(p, lab, bold=True, color=NAVY, size=BODY_SZ)
                segs[0] = ("t", rest)
        render_segments(p, segs, base_size=BODY_SZ)
    else:
        lab, rest = _lead_label(text, lead)
        if lab:
            add_run(p, lab, bold=True, color=NAVY, size=BODY_SZ)
            text = rest
        add_run(p, text, size=BODY_SZ)


def grades_table(doc, rows):
    widths = [2.60, 1.10]
    t = fixed_table(doc, widths, rows=len(rows) + 1)
    hdr = t.rows[0]
    for i, htxt in enumerate(("Component", "Weight")):
        c = hdr.cells[i]
        shade_cell(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cp(c)
        if i:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, htxt, bold=True, color="FFFFFF", size=11)
    for ri, (label, weight) in enumerate(rows, start=1):
        row = t.rows[ri]
        fill = "FFFFFF" if ri % 2 else CREAM
        for ci, val in enumerate((label, weight)):
            c = row.cells[ci]
            shade_cell(c, fill)
            cell_borders(c, top=(4, LIGHT), bottom=(4, LIGHT),
                         left=(4, LIGHT), right=(4, LIGHT))
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cp(c)
            if ci:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, bold=bool(ci), color=NAVY, size=11)
    spacer(doc, 6)


def title_block(doc):
    p = doc.add_paragraph()
    add_run(p, TITLE, bold=True, color=NAVY, size=20)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, SUBTITLE, bold=True, color=NAVY, size=13)
    gold_rule(doc, sz=20, space_after=5)

    for label, value in (("Instructor:", INSTRUCTOR + "  ·  " + OFFICE),
                         ("TA:", C.TA_NAME),
                         ("Classroom:", C.CLASSROOM),
                         ("Class times:", C.CLASS_TIMES)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_run(p, f"{label}  ", bold=True, color=NAVY, size=12)
        add_run(p, value, size=12)


def website_card(doc):
    """The page-1 callout, and the most prominent thing on the page after the
    title (2026-09-04, Nico): the website replaces last year's "see the
    separate section-specific calendar on Bruin Learn" line, and the PDF of
    the calendar is offered as a download FROM the website rather than as a
    reference of its own."""
    spacer(doc, 5)

    def pop(cell, inner_w):
        p = cp(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "Course Website", bold=True, color=NAVY, size=15)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        add_hyperlink(p, C.LINKS["website"], C.WEBSITE_TEXT, size=17,
                      underline=True)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "Everything for this class in one place: the schedule week "
                   "by week, all videos and podcasts, the readings, the "
                   "practice exercises and every deadline.",
                color=NAVY, size=11.5)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "Prefer a PDF? The course calendar and this syllabus can "
                   "be downloaded from the website.",
                italic=True, color=GRAY, size=10.5)

    # gold border at 2 pt, twice the weight of an ordinary card, so the box
    # carries the page
    rounded_card(doc, pop, fill=CREAM, border=GOLD, border_w=25400)


def build(doc):
    title_block(doc)
    website_card(doc)
    for blk in SECTIONS:
        kind = blk[0]
        if kind == "h":
            heading(doc, blk[1])
        elif kind == "sub":
            sub_heading(doc, blk[1])
        elif kind == "p":
            body(doc, text=blk[1])
        elif kind == "mix":
            body(doc, segs=blk[1])
        elif kind == "b":
            bullet(doc, text=blk[1])
        elif kind == "bmix":
            bullet(doc, segs=blk[1])
        elif kind == "bl":
            bullet(doc, text=blk[1], lead=True)
        elif kind == "blmix":
            bullet(doc, segs=blk[1], lead=True)
        elif kind == "grades":
            grades_table(doc, blk[1])
        elif kind == "gap":
            spacer(doc, blk[1])
        else:
            raise ValueError("unknown block kind %r" % (kind,))


# ============================== Markdown draft ==============================

def _md_lead(text, lead):
    if not lead:
        return text
    i = text.find(":")
    return "**%s**%s" % (text[:i + 1], text[i + 1:])


def md_segments(segs):
    out = []
    for seg in segs:
        if seg[0] == "t":
            out.append(seg[1])
        else:
            out.append("[%s](%s)" % (seg[2], C.LINKS[seg[1]]))
    return "".join(out)


def write_md():
    L = ["# %s" % TITLE, "", "## %s" % SUBTITLE, "",
         "**Instructor:** %s  ·  %s  " % (INSTRUCTOR, OFFICE),
         "**TA:** %s  " % C.TA_NAME,
         "**Classroom:** %s  " % C.CLASSROOM,
         "**Class times:** %s" % C.CLASS_TIMES, "",
         "> ### Course Website",
         "> **[%s](%s)**" % (C.WEBSITE_TEXT, C.LINKS["website"]),
         ">",
         "> Everything for this class in one place: the schedule week by "
         "week, all videos and podcasts, the readings, the practice "
         "exercises and every deadline.",
         ">",
         "> *Prefer a PDF? The course calendar and this syllabus can be "
         "downloaded from the website.*", ""]
    for blk in SECTIONS:
        k = blk[0]
        if k == "h":
            L += ["## %s" % blk[1], ""]
        elif k == "sub":
            L += ["### %s" % blk[1], ""]
        elif k == "p":
            L += [blk[1], ""]
        elif k == "mix":
            L += [md_segments(blk[1]), ""]
        elif k in ("b", "bl"):
            L += ["- %s" % _md_lead(blk[1], k == "bl")]
        elif k in ("bmix", "blmix"):
            L += ["- %s" % _md_lead(md_segments(blk[1]), k == "blmix")]
        elif k == "grades":
            L += ["", "| Component | Weight |", "|---|---|"]
            L += ["| %s | %s |" % r for r in blk[1]]
            L += [""]
        if k in ("b", "bl", "bmix", "blmix"):
            continue
    # a blank line after every bullet run
    out = []
    for i, line in enumerate(L):
        out.append(line)
        if line.startswith(("- ", "- **")) and (
                i + 1 >= len(L) or not L[i + 1].startswith("- ")):
            out.append("")
    io.open(OUT_MD, "w", encoding="utf-8", newline="\n").write(
        "\n".join(out).rstrip() + "\n")
    print("Saved: %s" % OUT_MD)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--section", default=os.environ.get("MGMT405_SECTION"),
                    help="emba (default) or femba; read before the imports")
    ap.add_argument("--md", action="store_true",
                    help="also write the Markdown draft")
    args = ap.parse_args()

    doc = setup_document()
    build(doc)
    doc.save(OUT)
    print("Saved: %s" % OUT)
    if args.md:
        write_md()


if __name__ == "__main__":
    main()
