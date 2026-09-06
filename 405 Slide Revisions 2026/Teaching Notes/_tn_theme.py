"""Shared Word theme layer for the MGMT 405 teaching notes.

One place for the course palette, the masthead, the section headings, the
cream callout card, the native-figure primitives and the tracked-change
helpers.  Each note has its own `_build_TN_*.py` that imports this module.

The visual reference is the course calendar
(`405 Calendar and Website/Course Calendar/_build_calendar.py`): navy /
gold / gray on white, Calibri throughout, a gold rule under the masthead
and rounded cream cards for asides.

Tracked changes: wording edits are emitted as real Word revisions
(`w:ins` / `w:del`) so Nico can accept or reject each one in the Review
pane.  Formatting is NOT tracked -- a restyle that marked every paragraph
as changed would bury the wording edits.
"""

import itertools

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

# python-docx does not ship the 2010 drawing namespaces; register them so
# nsdecls() can emit them on the fragments we parse in.
from docx.oxml.ns import nsmap as _nsmap
_nsmap.setdefault(
    "wps", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")
_nsmap.setdefault(
    "wpg", "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup")

# --------------------------------------------------------------------------
# Palette (Teaching CLAUDE.md, "Canvas, palette, and chrome")
# --------------------------------------------------------------------------
NAVY     = "0B2B4E"   # primary: headings, axes, structural lines
GOLD     = "E09F3E"   # accent: rules, callout borders
GRAY     = "555B66"   # neutral: captions, secondary text
LIGHT    = "C8CDD3"   # thin rules
CREAM    = "FDF6E6"   # callout / card fill
PALEGOLD = "F6E8C9"
DARKRED  = "C00000"   # demand curves (2026-08-30 rule)
CBLUE    = "0070C0"   # reserved concept blue: MR, concept names
WHITE    = "FFFFFF"

EMU_IN = 914400


def emu(inches):
    return int(round(inches * EMU_IN))


# --------------------------------------------------------------------------
# Revision (tracked-change) bookkeeping
# --------------------------------------------------------------------------
REV_AUTHOR = "Claude (proposed)"
REV_DATE = "2026-09-06T00:00:00Z"
_rev_ids = itertools.count(900)


def _stamp(el):
    """Give a w:ins / w:del element its id / author / date."""
    el.set(qn("w:id"), str(next(_rev_ids)))
    el.set(qn("w:author"), REV_AUTHOR)
    el.set(qn("w:date"), REV_DATE)
    return el


# --------------------------------------------------------------------------
# Runs
# --------------------------------------------------------------------------
def style_run(r, bold=False, italic=False, underline=False, color=None,
              size=None, font="Calibri", subscript=False, superscript=False):
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = font
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = RGBColor.from_string(color)
    if subscript:
        r.font.subscript = True
    if superscript:
        r.font.superscript = True
    # East-Asian / complex-script font, so Word does not substitute
    rPr = r._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), font)
    rf.set(qn("w:cs"), font)
    return r


def run(p, text, **kw):
    """Plain (untracked) run."""
    return style_run(p.add_run(text), **kw)


def ins(p, text, **kw):
    """Run marked as a tracked INSERTION."""
    w = _stamp(OxmlElement("w:ins"))
    p._p.append(w)
    r = style_run(p.add_run(text), **kw)
    w.append(r._r)
    return r


def dele(p, text, **kw):
    """Run marked as a tracked DELETION (w:t becomes w:delText)."""
    w = _stamp(OxmlElement("w:del"))
    p._p.append(w)
    r = style_run(p.add_run(text), **kw)
    for t in r._r.findall(qn("w:t")):
        t.tag = qn("w:delText")
    w.append(r._r)
    return r


def para_inserted(p):
    """Mark the paragraph MARK as inserted -- used for wholly new paragraphs.

    Call this AFTER the paragraph's spacing / alignment are set: w:rPr sits
    near the end of w:pPr in schema order.
    """
    pPr = p._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    rPr.insert(0, _stamp(OxmlElement("w:ins")))
    return p


# --------------------------------------------------------------------------
# Paragraphs
# --------------------------------------------------------------------------
def para(doc, before=0, after=8, align=None, line=None, left=0.0, hang=0.0,
         keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if align is not None:
        pf.alignment = align
    if line is not None:
        pf.line_spacing = line
    if left:
        pf.left_indent = Inches(left)
    if hang:
        pf.first_line_indent = Inches(-hang)
    if keep_next:
        p.paragraph_format.keep_with_next = True
    return p


def body(doc, text=None, **kw):
    """Justified body paragraph, matching the original notes."""
    kw.setdefault("after", 9)
    kw.setdefault("align", WD_ALIGN_PARAGRAPH.JUSTIFY)
    p = para(doc, **kw)
    if text:
        run(p, text, size=11)
    return p


def heading(doc, text, before=14, after=7):
    """Navy bold section heading (replaces the original's black underline)."""
    p = para(doc, before=before, after=after, keep_next=True)
    run(p, text, bold=True, color=NAVY, size=13)
    return p


def gold_rule(p, size_eighths=18, color=GOLD, space=1):
    """Put a horizontal rule under a paragraph (used beneath the masthead)."""
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    bdr.append(bottom)
    pPr.append(bdr)
    return p


# --------------------------------------------------------------------------
# Document skeleton
# --------------------------------------------------------------------------
COURSE_TITLE = "MGMT 405 \u2013 Managerial Economics"
FOOTER_LEFT = "MGMT 405  \u00b7  Fall 2026"


def new_doc(margin_in=1.0):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    rPr = normal.element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), "Calibri")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.0

    sec = doc.sections[0]
    sec.top_margin = Inches(margin_in)
    sec.bottom_margin = Inches(margin_in)
    sec.left_margin = Inches(margin_in)
    sec.right_margin = Inches(margin_in)

    return doc


def footer(doc, left_text=None):
    """Gray footer with a live page-number field.

    With `left_text`, that sits at the left and the number at the right.
    Without it (the teaching-note default), a lone centred page number: the
    notes carry no year or term, so the same file is reusable next year
    without being rebuilt (2026-09-06, Nico).
    """
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if left_text:
        width = sec.page_width - sec.left_margin - sec.right_margin
        p.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
        # Two tabs: the Footer style already carries a centre stop at
        # mid-column, so one tab parks the page number in the middle.
        run(p, left_text + "\t\t", color=GRAY, size=9)
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = style_run(p.add_run(), color=GRAY, size=9)
    _page_field(r)
    return p


def _page_field(r):
    """Live PAGE field, so the number never drifts."""
    for kind, txt in (("begin", None), (None, "PAGE"), ("separate", None)):
        if kind:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            r._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = " PAGE "
            r._r.append(it)
    t = OxmlElement("w:t")
    t.text = "1"
    r._r.append(t)
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r._r.append(fc)
    return r


def callout(doc, prefix, text, width_in=6.5, size=10.5, fill=CREAM,
            before=6, after=6):
    """The house cream aside: bold navy prefix, then the body text.

    Used for the "Convention:" / "Recall:" / "Note:" boxes.  The originals
    carried some of these as Word footnotes; a cream card is the deck-side
    device for the same job and puts the aside where a student sees it.

    `text` is either a plain string or a list of (text, props) run specs,
    for an aside that carries its own emphasis inside the box.
    """
    def pop(cell):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if prefix:
            run(p, prefix + "  ", bold=True, color=NAVY, size=size)
        if isinstance(text, str):
            run(p, text, color=NAVY, size=size)
        else:
            for t, props in text:
                kw = dict(color=NAVY, size=size)
                kw.update(props)
                run(p, t, **kw)

    return cream_card(doc, pop, fill=fill, width_in=width_in,
                      before=before, after=after)


def masthead(doc, subtitle):
    """Calendar-style masthead: navy course line, navy subtitle, gold rule."""
    p = para(doc, before=0, after=1)
    run(p, COURSE_TITLE, bold=True, color=NAVY, size=20)

    p2 = para(doc, before=0, after=9)
    run(p2, subtitle, bold=True, color=NAVY, size=13)
    gold_rule(p2, size_eighths=18, space=6)
    return p2


# --------------------------------------------------------------------------
# Rounded cream card (the "Convention" callout, Word edition)
# --------------------------------------------------------------------------
_card_id = [500]


def cream_card(doc, populate, fill=CREAM, border=NAVY, width_in=6.5,
               border_w=9525, height_in=None, before=6, after=6):
    """A rounded, shaded, drop-shadowed card holding `populate(cell)`'s text.

    Built the way the course calendar builds its cards: render the
    paragraphs into a throwaway table cell, measure them, then transplant
    them into a real roundRect shape's text box and drop the table.  The
    shape carries the fill, the border and the shadow, so the corners are
    genuinely rounded rather than a floating outline that can drift out of
    register with the text.
    """
    inner_w = width_in - 0.24               # 0.12" text inset each side
    tmp = doc.add_table(rows=1, cols=1)
    tmp.autofit = False
    cell = tmp.cell(0, 0)
    cell.width = Inches(inner_w)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    populate(cell)

    pars = [ch for ch in cell._tc if ch.tag == qn("w:p")]
    h_in = height_in or (sum(_measure_par(p, inner_w) for p in pars) + 0.16)

    cx, cy = emu(width_in), emu(h_in)
    adj = int(min(18000, max(3000, 0.11 / h_in * 100000)))
    _card_id[0] += 1
    drawing = parse_xml(_CARD_XML.format(
        nsd=nsdecls("w", "wp", "a", "wps"), cx=cx, cy=cy,
        did=_card_id[0], adj=adj, fill=fill, border=border, w=border_w))
    txbx = drawing.find(".//" + qn("w:txbxContent"))
    for p_el in pars:
        txbx.append(p_el)

    holder = para(doc, before=before, after=after)
    holder.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.add_run()._r.append(drawing)
    tmp._tbl.getparent().remove(tmp._tbl)
    return h_in


_CARD_XML = (
    '<w:drawing {nsd}>'
    '<wp:inline distT="0" distB="0" distL="0" distR="0">'
    '<wp:extent cx="{cx}" cy="{cy}"/>'
    '<wp:effectExtent l="0" t="0" r="76200" b="76200"/>'
    '<wp:docPr id="{did}" name="Card {did}"/>'
    '<a:graphic><a:graphicData'
    ' uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<wps:wsp><wps:cNvSpPr/><wps:spPr bwMode="auto">'
    '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
    '<a:prstGeom prst="roundRect"><a:avLst>'
    '<a:gd name="adj" fmla="val {adj}"/></a:avLst></a:prstGeom>'
    '<a:solidFill><a:srgbClr val="{fill}"/></a:solidFill>'
    '<a:ln w="{w}"><a:solidFill><a:srgbClr val="{border}"/></a:solidFill></a:ln>'
    '<a:effectLst><a:outerShdw blurRad="63500" dist="27940" dir="5400000"'
    ' rotWithShape="0"><a:srgbClr val="000000"><a:alpha val="30000"/>'
    '</a:srgbClr></a:outerShdw></a:effectLst>'
    '</wps:spPr><wps:txbx><w:txbxContent/></wps:txbx>'
    '<wps:bodyPr rot="0" vert="horz" wrap="square" lIns="109728" tIns="54864"'
    ' rIns="109728" bIns="54864" anchor="t" anchorCtr="0">'
    '<a:noAutofit/></wps:bodyPr></wps:wsp>'
    '</a:graphicData></a:graphic></wp:inline></w:drawing>'
)


def _measure_par(p_el, width_in):
    """Rendered height in inches of one w:p laid out at `width_in`.

    Measured RUN BY RUN in each run's own face: bold Calibri is noticeably
    wider than regular, and measuring a mostly-bold paragraph in the regular
    face under-counts the lines and clips the card (2026-09-06).  OMML runs
    count double in width and get a taller line, so a stacked fraction sizes
    a hero-formula card correctly.
    """
    import math

    width = 0.0
    size, has_math, any_text = 10.5, False, False

    for r in p_el.iter():
        if r.tag not in (qn("w:r"), qn("m:r")):
            continue
        txt = "".join(t.text or "" for t in r.iter()
                      if t.tag in (qn("w:t"), qn("w:delText"), qn("m:t")))
        if not txt:
            continue
        any_text = True
        is_math = r.tag == qn("m:r") or r.find(".//" + qn("m:t")) is not None
        has_math = has_math or is_math

        rPr = r.find(qn("w:rPr"))
        bold = italic = False
        rsize = 11.0
        if rPr is not None:
            bold = rPr.find(qn("w:b")) is not None
            italic = rPr.find(qn("w:i")) is not None
            szel = rPr.find(qn("w:sz"))
            if szel is not None:
                try:
                    rsize = int(szel.get(qn("w:val"))) / 2.0
                except (TypeError, ValueError):
                    pass
        size = max(size, rsize)
        w = _font(rsize, bold, italic).getlength(txt) / (8 * 72.0)
        width += w * (2.0 if is_math else 1.0)

    if not any_text:
        return 0.10

    avail = max(width_in, 0.5)
    # 1.04 of slack: justified text breaks a little earlier than the raw
    # glyph sum implies, and a clipped card is far worse than a roomy one.
    lines = max(1, int(math.ceil(width * 1.04 / avail)))
    line_h = size * (2.35 if has_math else 1.34) / 72.0

    extra = 0.0
    pPr = p_el.find(qn("w:pPr"))
    if pPr is not None:
        sp = pPr.find(qn("w:spacing"))
        if sp is not None:
            for a in ("w:before", "w:after"):
                v = sp.get(qn(a))
                if v:
                    extra += int(v) / 20.0 / 72.0
    return lines * line_h + extra


def _cell_margins(cell, pad_in):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, v in (("top", pad_in), ("start", pad_in + 0.04),
                    ("bottom", pad_in), ("end", pad_in + 0.04)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(int(v * 1440)))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _rounded_outline(cell, color, w, width_in):
    """Overlay a rounded outline so the shaded cell reads as a lifted card."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    r = p.add_run()
    r._r.append(parse_xml(_ROUNDED_XML.format(
        nsd=nsdecls("w", "wp", "a", "wps"),
        cx=emu(width_in + 0.06), cy=emu(0.55),
        color=color, w=w)))


_ROUNDED_XML = (
    '<w:drawing {nsd}>'
    '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
    ' relativeHeight="2" behindDoc="1" locked="0" layoutInCell="1"'
    ' allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>-38100</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>-38100</wp:posOffset></wp:positionV>'
    '<wp:extent cx="{cx}" cy="{cy}"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
    '<wp:wrapNone/><wp:docPr id="0" name=""/><wp:cNvGraphicFramePr/>'
    '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<wps:wsp><wps:cNvSpPr/>'
    '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
    '<a:prstGeom prst="roundRect"><a:avLst>'
    '<a:gd name="adj" fmla="val 12000"/></a:avLst></a:prstGeom>'
    '<a:noFill/>'
    '<a:ln w="{w}"><a:solidFill><a:srgbClr val="{color}"/></a:solidFill></a:ln>'
    '</wps:spPr><wps:bodyPr/></wps:wsp>'
    '</a:graphicData></a:graphic></wp:anchor></w:drawing>'
)


# --------------------------------------------------------------------------
# Native tables -- navy header, cream/white body rows, thin borders
# --------------------------------------------------------------------------
def table(doc, rows, widths_in, header=True, size=9.5, highlight=(),
          align_right=(), italic_rows=(), label_rows=(), keep_together=True,
          replaced=None):
    """A course-styled table.

    `rows`     list of row lists; cell text as plain strings.
    `widths_in` column widths in inches.
    `highlight` (row, col) pairs to mark with the house pale-gold cell fill
                and bold navy text -- the deck's "table-cell number
                highlight" device (kept flat, no rounding, no shadow).
    `align_right` column indices to right-align.
    `italic_rows` row indices set italic (Excel's sub-header rows).
    `label_rows`  row indices that are section labels (ANOVA, and the like):
                  navy bold, no cream banding.
    """
    tbl = doc.add_table(rows=len(rows), cols=len(widths_in))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _strip_table_style(tbl)

    for c, w in enumerate(widths_in):
        tbl.columns[c].width = Inches(w)

    hl = set(highlight)
    for r, row in enumerate(rows):
        for c in range(len(widths_in)):
            cell = tbl.cell(r, c)
            cell.width = Inches(widths_in[c])
            _cell_margins(cell, 0.045)
            txt = row[c] if c < len(row) else ""

            is_head = header and r == 0
            is_hl = (r, c) in hl
            if is_head:
                _shade(cell, NAVY)
            elif is_hl:
                _shade(cell, PALEGOLD)
            elif r in label_rows:
                _shade(cell, WHITE)
            else:
                _shade(cell, CREAM if r % 2 == 0 else WHITE)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if c in align_right:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            props = dict(bold=is_head or is_hl or (r in label_rows),
                         italic=(r in italic_rows),
                         color=WHITE if is_head else NAVY,
                         size=size)
            swap = (replaced or {}).get((r, c))
            if swap:
                # a tracked cell replacement: old struck through, new inserted
                dele(p, swap[0], **props)
                ins(p, swap[1], **props)
            elif txt:
                run(p, txt, **props)

    _thin_borders(tbl)
    if keep_together:
        _keep_table_together(tbl)
    return tbl


def _keep_table_together(tbl):
    """Stop Word breaking the table across a page.

    Two parts: no row may split internally (w:cantSplit), and every
    paragraph outside the last row keeps with the next, so Word moves the
    whole block down rather than orphaning the first rows.
    """
    rows = tbl.rows
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if i < len(rows) - 1:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True


def _strip_table_style(tbl):
    """Drop python-docx's default table style; fills are set explicitly."""
    tblPr = tbl._tbl.tblPr
    for st in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(st)
    look = tblPr.find(qn("w:tblLook"))
    if look is not None:
        tblPr.remove(look)


def _thin_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), LIGHT)
        borders.append(e)
    tblPr.append(borders)


# --------------------------------------------------------------------------
# Pictures -- rounded corners plus a soft shadow, so they read as lifted
# --------------------------------------------------------------------------
def picture(doc, path, width_in, caption=None, rounded=True, shadow=True,
            before=8, after=4):
    """Place an image, optionally rounded / shadowed, with a gray caption.

    Kept as an image on purpose where the underlying data is not available
    to rebuild the chart natively -- fabricating data points would be worse
    than reusing Nico's own figure.
    """
    p = para(doc, before=before, after=0 if caption else after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width_in))

    if rounded or shadow:
        _lift_picture(r, rounded, shadow)

    if caption:
        cp = para(doc, before=3, after=after)
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp, caption, italic=True, color=GRAY, size=9)
    return p


def _lift_picture(r, rounded, shadow):
    """Add roundRect geometry and an outer shadow to the last inline picture."""
    spPr = r._r.find(".//" + qn("a:graphic") + "/" + qn("a:graphicData"))
    if spPr is None:
        return
    pic = spPr.find(qn("pic:pic"))
    if pic is None:
        return
    sp = pic.find(qn("pic:spPr"))
    if sp is None:
        return
    if rounded:
        for g in sp.findall(qn("a:prstGeom")):
            sp.remove(g)
        geom = parse_xml(
            '<a:prstGeom {ns} prst="roundRect"><a:avLst>'
            '<a:gd name="adj" fmla="val 2600"/></a:avLst></a:prstGeom>'
            .format(ns=nsdecls("a")))
        xfrm = sp.find(qn("a:xfrm"))
        sp.insert(list(sp).index(xfrm) + 1 if xfrm is not None else 0, geom)
    if shadow:
        sp.append(parse_xml(
            '<a:effectLst {ns}><a:outerShdw blurRad="50800" dist="25400"'
            ' dir="2700000" algn="tl" rotWithShape="0">'
            '<a:srgbClr val="000000"><a:alpha val="30000"/></a:srgbClr>'
            '</a:outerShdw></a:effectLst>'.format(ns=nsdecls("a"))))


# --------------------------------------------------------------------------
# Native figures (DrawingML shape groups -- never screenshots)
# --------------------------------------------------------------------------
class Fig:
    """A DrawingML shape group, addressed in inches from its top-left.

    Mirrors the deck-side `Fig` helper: a small set of primitives that emit
    real, editable Word shapes so a figure can be nudged in Word afterward.
    """

    def __init__(self, w_in, h_in, name="Figure"):
        self.w, self.h = w_in, h_in
        self.name = name
        self.shapes = []
        self._id = itertools.count(2)

    # -- primitives --------------------------------------------------------
    def line(self, x1, y1, x2, y2, color=NAVY, w_pt=1.25, dash=None,
             arrow=False, name="line"):
        flipH = x2 < x1
        flipV = y2 < y1
        x, y = min(x1, x2), min(y1, y2)
        cx, cy = abs(x2 - x1), abs(y2 - y1)
        xfrm = '<a:xfrm{f}><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'.format(
            f=(' flipH="1"' if flipH else "") + (' flipV="1"' if flipV else ""),
            x=emu(x), y=emu(y), cx=max(emu(cx), 1), cy=max(emu(cy), 1))
        ln = '<a:ln w="{w}" cap="rnd"><a:solidFill><a:srgbClr val="{c}"/></a:solidFill>'.format(
            w=int(w_pt * 12700), c=color)
        if dash:
            ln += '<a:prstDash val="{}"/>'.format(dash)
        ln += "<a:round/>"
        if arrow:
            ln += '<a:tailEnd type="triangle" w="med" len="med"/>'
        ln += "</a:ln>"
        self.shapes.append(
            '<wps:wsp><wps:cNvPr id="{i}" name="{n}"/><wps:cNvSpPr/>'
            '<wps:spPr>{xfrm}<a:prstGeom prst="line"><a:avLst/></a:prstGeom>'
            '<a:noFill/>{ln}</wps:spPr>'
            '<wps:bodyPr/></wps:wsp>'.format(i=next(self._id), n=name,
                                             xfrm=xfrm, ln=ln))

    def curve(self, pts, color=NAVY, w_pt=1.75, name="curve"):
        """One editable freeform through quadratic Bezier anchors.

        `pts` is [start, (ctrl, end), (ctrl, end), ...] in figure inches, so
        "Edit Points" in Word shows a handful of handles rather than a dense
        polyline.  Element order inside <a:ln> matters: fill, then dash,
        then join -- and <a:custGeom> sits after <a:xfrm>, before the fill.
        """
        xs = [pts[0][0]] + [c for seg in pts[1:] for c in (seg[0][0], seg[1][0])]
        ys = [pts[0][1]] + [c for seg in pts[1:] for c in (seg[0][1], seg[1][1])]
        x0, y0 = min(xs), min(ys)
        cw, ch = max(max(xs) - x0, 0.01), max(max(ys) - y0, 0.01)
        W, H = emu(cw), emu(ch)

        def pt(p):
            return ('<a:pt x="{x}" y="{y}"/>'
                    .format(x=int(round((p[0] - x0) / cw * W)),
                            y=int(round((p[1] - y0) / ch * H))))

        path = "<a:moveTo>{}</a:moveTo>".format(pt(pts[0]))
        for ctrl, end in pts[1:]:
            path += "<a:quadBezTo>{}{}</a:quadBezTo>".format(pt(ctrl), pt(end))

        self.shapes.append(
            '<wps:wsp><wps:cNvPr id="{i}" name="{n}"/><wps:cNvSpPr/>'
            '<wps:spPr>'
            '<a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{W}" cy="{H}"/></a:xfrm>'
            '<a:custGeom><a:avLst/><a:gdLst/><a:ahLst/><a:cxnLst/>'
            '<a:rect l="0" t="0" r="{W}" b="{H}"/>'
            '<a:pathLst><a:path w="{W}" h="{H}">{path}</a:path></a:pathLst>'
            '</a:custGeom>'
            '<a:noFill/>'
            '<a:ln w="{lw}" cap="rnd"><a:solidFill><a:srgbClr val="{c}"/>'
            '</a:solidFill><a:round/></a:ln>'
            '</wps:spPr><wps:bodyPr/></wps:wsp>'.format(
                i=next(self._id), n=name, x=emu(x0), y=emu(y0), W=W, H=H,
                path=path, lw=int(w_pt * 12700), c=color))

    def dot(self, x, y, d=0.075, color=NAVY, name="dot"):
        self.shapes.append(
            '<wps:wsp><wps:cNvPr id="{i}" name="{n}"/><wps:cNvSpPr/>'
            '<wps:spPr><a:xfrm><a:off x="{x}" y="{y}"/>'
            '<a:ext cx="{d}" cy="{d}"/></a:xfrm>'
            '<a:prstGeom prst="ellipse"><a:avLst/></a:prstGeom>'
            '<a:solidFill><a:srgbClr val="{c}"/></a:solidFill>'
            '<a:ln><a:noFill/></a:ln></wps:spPr>'
            '<wps:bodyPr/></wps:wsp>'.format(
                i=next(self._id), n=name, x=emu(x - d / 2), y=emu(y - d / 2),
                d=emu(d), c=color))

    def label(self, x, y, runs, w=None, h=0.24, size=11, align="l",
              fill=None, border=None, name="lbl", pad=0.0):
        """A text label.  `runs` is a list of (text, dict-of-run-props).

        `x, y` is the top-left of the box, unless align is "c" / "r", in
        which case x is the box's horizontal centre / right edge.
        """
        if w is None:
            w = text_width(runs) + 2 * pad + 0.06
        if align == "c":
            x = x - w / 2
        elif align == "r":
            x = x - w
        body_runs = "".join(_wrun(t, pr) for t, pr in runs)
        jc = {"l": "left", "c": "center", "r": "right"}[align]
        geom = ('<a:prstGeom prst="roundRect"><a:avLst>'
                '<a:gd name="adj" fmla="val {adj}"/></a:avLst></a:prstGeom>'
                .format(adj=int(0.06 / min(w, h) * 100000))
                if border or fill else
                '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>')
        fillxml = ('<a:solidFill><a:srgbClr val="{}"/></a:solidFill>'.format(fill)
                   if fill else "<a:noFill/>")
        lnxml = ('<a:ln w="9525"><a:solidFill><a:srgbClr val="{}"/></a:solidFill></a:ln>'
                 .format(border) if border else "<a:ln><a:noFill/></a:ln>")
        ins_ = emu(pad)
        self.shapes.append(
            '<wps:wsp><wps:cNvPr id="{i}" name="{n}"/><wps:cNvSpPr txBox="1"/>'
            '<wps:spPr><a:xfrm><a:off x="{x}" y="{y}"/>'
            '<a:ext cx="{cx}" cy="{cy}"/></a:xfrm>{geom}{fill}{ln}</wps:spPr>'
            '<wps:txbx><w:txbxContent><w:p><w:pPr>'
            '<w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
            '<w:jc w:val="{jc}"/></w:pPr>{runs}</w:p></w:txbxContent></wps:txbx>'
            '<wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"'
            ' horzOverflow="overflow" vert="horz" wrap="square"'
            ' lIns="{p}" tIns="{p}" rIns="{p}" bIns="{p}" anchor="ctr"'
            ' anchorCtr="0"><a:noAutofit/></wps:bodyPr></wps:wsp>'.format(
                i=next(self._id), n=name, x=emu(x), y=emu(y),
                cx=emu(w), cy=emu(h), geom=geom, fill=fillxml, ln=lnxml,
                jc=jc, runs=body_runs, p=ins_))

    # -- assembly ----------------------------------------------------------
    def xml(self, doc_pr_id=101, scale=1.0):
        return (
            '<w:drawing {nsd}>'
            '<wp:inline distT="0" distB="0" distL="0" distR="0">'
            '<wp:extent cx="{cx}" cy="{cy}"/>'
            '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            '<wp:docPr id="{did}" name="{name}"/>'
            '<wp:cNvGraphicFramePr/>'
            '<a:graphic><a:graphicData'
            ' uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">'
            '<wpg:wgp><wpg:cNvGrpSpPr/><wpg:grpSpPr><a:xfrm>'
            '<a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/>'
            '<a:chOff x="0" y="0"/><a:chExt cx="{chcx}" cy="{chcy}"/>'
            '</a:xfrm></wpg:grpSpPr>{shapes}</wpg:wgp>'
            '</a:graphicData></a:graphic></wp:inline></w:drawing>'
        ).format(nsd=nsdecls("w", "wp", "a", "wps", "wpg"),
                 cx=emu(self.w * scale), cy=emu(self.h * scale),
                 chcx=emu(self.w), chcy=emu(self.h),
                 did=doc_pr_id, name=self.name,
                 shapes="".join(self.shapes))

    def place(self, doc, before=6, after=6, scale=1.0):
        p = para(doc, before=before, after=after)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run()._r.append(parse_xml(self.xml(scale=scale)))
        return p


def _wrun(text, pr):
    """One w:r inside a txbxContent paragraph."""
    props = ['<w:rFonts w:ascii="{f}" w:hAnsi="{f}" w:cs="{f}"/>'.format(
        f=pr.get("font", "Calibri"))]
    if pr.get("bold"):
        props.append("<w:b/>")
    if pr.get("italic"):
        props.append("<w:i/>")
    if pr.get("color"):
        props.append('<w:color w:val="{}"/>'.format(pr["color"]))
    if pr.get("subscript"):
        props.append('<w:vertAlign w:val="subscript"/>')
    if pr.get("superscript"):
        props.append('<w:vertAlign w:val="superscript"/>')
    sz = int(pr.get("size", 11) * 2)
    props.append('<w:sz w:val="{s}"/><w:szCs w:val="{s}"/>'.format(s=sz))
    esc = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    return ('<w:r><w:rPr>{p}</w:rPr><w:t xml:space="preserve">{t}</w:t></w:r>'
            .format(p="".join(props), t=esc))


# --------------------------------------------------------------------------
# Text measurement -- boxes are sized to their label, never guessed
# --------------------------------------------------------------------------
_FONT_CACHE = {}


def _font(size_pt, bold=False, italic=False):
    from PIL import ImageFont
    key = (size_pt, bold, italic)
    if key in _FONT_CACHE:
        return _FONT_CACHE[key]
    name = {(False, False): "calibri.ttf", (True, False): "calibrib.ttf",
            (False, True): "calibrii.ttf", (True, True): "calibriz.ttf"}[(bold, italic)]
    for base in (r"C:\Windows\Fonts", r"C:\Windows\Fonts\\"):
        try:
            f = ImageFont.truetype(base + "\\" + name, int(round(size_pt * 8)))
            _FONT_CACHE[key] = f
            return f
        except OSError:
            continue
    raise OSError("Calibri not found: " + name)


def text_width(runs):
    """Rendered width in inches of a list of (text, props) runs."""
    total = 0.0
    for t, pr in runs:
        size = pr.get("size", 11)
        if pr.get("subscript") or pr.get("superscript"):
            size *= 0.65
        f = _font(size, pr.get("bold", False), pr.get("italic", False))
        total += f.getlength(t) / (8 * 72.0)
    return total


# --------------------------------------------------------------------------
# OMML helpers
# --------------------------------------------------------------------------
MATH_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def mrun(text, italic=None):
    """One OMML run.  Cambria Math; `italic=False` sets an upright acronym."""
    sty = ""
    if italic is False:
        sty = '<m:rPr><m:sty m:val="p"/></m:rPr>'
    esc = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return ('<m:r>{s}<w:rPr><w:rFonts w:ascii="Cambria Math"'
            ' w:hAnsi="Cambria Math"/></w:rPr><m:t xml:space="preserve">{t}</m:t></m:r>'
            .format(s=sty, t=esc))


def mfrac(num, den):
    ctrl = ('<m:fPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math"'
            ' w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:fPr>')
    return "<m:f>{c}<m:num>{n}</m:num><m:den>{d}</m:den></m:f>".format(
        c=ctrl, n=num, d=den)


def msub(base, sub):
    """Subscripted symbol, e.g. msub(mrun("MRP", italic=False), mrun("L"))."""
    ctrl = ('<m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math"'
            ' w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>')
    return "<m:sSub>{c}<m:e>{b}</m:e><m:sub>{s}</m:sub></m:sSub>".format(
        c=ctrl, b=base, s=sub)


def acr(name, sub=None):
    """A multi-letter acronym set upright, with an optional italic subscript.

    Course convention: acronyms upright (m:sty p), single-letter variables
    italic.  So MRP_L is upright "MRP" with an italic "L" subscript.
    """
    base = mrun(name, italic=False)
    return msub(base, mrun(sub)) if sub else base


def msup(base, sup):
    ctrl = ('<m:sSupPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math"'
            ' w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSupPr>')
    return "<m:sSup>{c}<m:e>{b}</m:e><m:sup>{s}</m:sup></m:sSup>".format(
        c=ctrl, b=base, s=sup)


def m_ins(inner):
    """Wrap OMML content in a tracked insertion."""
    return '<w:ins w:id="{i}" w:author="{a}" w:date="{d}">{x}</w:ins>'.format(
        i=next(_rev_ids), a=REV_AUTHOR, d=REV_DATE, x=inner)


def m_del(inner):
    """Wrap OMML content in a tracked deletion (math keeps m:t, not delText)."""
    return '<w:del w:id="{i}" w:author="{a}" w:date="{d}">{x}</w:del>'.format(
        i=next(_rev_ids), a=REV_AUTHOR, d=REV_DATE, x=inner)


def _size_omml(content, size):
    """Force a point size on every run inside an OMML fragment."""
    if size is None:
        return content
    sz = '<w:sz w:val="{n}"/><w:szCs w:val="{n}"/>'.format(n=int(size * 2))
    return content.replace("</w:rPr>", sz + "</w:rPr>")


def color_omml(content, color):
    """Recolor every run in an OMML fragment.

    `w:color` has to sit after `w:rFonts` and before `w:sz` in schema
    order, so it is spliced in right behind the font element -- appended at
    the end of `w:rPr` it would be after `w:sz` and Word ignores it.
    """
    tag = '<w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>'
    return content.replace(
        tag, tag + '<w:color w:val="{}"/>'.format(color))


def equation(doc, content, before=4, after=10, size=None):
    """A centred display equation (m:oMathPara) in its own paragraph."""
    p = para(doc, before=before, after=after)
    xml = ('<m:oMathPara {m} {w}><m:oMath>{c}</m:oMath></m:oMathPara>'
           .format(m=MATH_NS, w=W_NS, c=_size_omml(content, size)))
    p._p.append(parse_xml(xml))
    return p


def equation_inline(p, content, size=None):
    """An OMML fragment inside a running paragraph (no oMathPara wrapper)."""
    xml = ('<m:oMath {m} {w}>{c}</m:oMath>'
           .format(m=MATH_NS, w=W_NS, c=_size_omml(content, size)))
    p._p.append(parse_xml(xml))
    return p
